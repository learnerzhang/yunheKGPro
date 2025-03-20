# coding=utf-8

from django.forms import model_to_dict
from pydantic import BaseModel
from typing import List
import os
import json
from docx.shared import Mm
import re
import base64
import jinja2
import time
import sys
import requests
from docx import Document
from datetime import datetime
import pandas as pd
from io import StringIO
from tabulate import tabulate
import requests
import json
import difflib

from userapp.models import User
from yaapp import divHtml, pd2HtmlCSS, text_table, extract_shuiku_data,extract_shuiku_data_jianyi,yujingdengji,search_fragpacks
from yaapp.models import PlanTemplate
from langchain.llms import Ollama
def get_access_token():
    """
    使用 API Key，Secret Key 获取access_token，替换下列示例中的应用API Key、应用Secret Key
    """
    #文 心千帆模型调用
    API_KEY = "cP3fnwcnK9MzGsWGfGEPeUkP"
    SECRET_KEY = "AfeuE0n4ZY5eCYG21QqSXG6WkBGltcxM"
    # url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=[应用API Key]&client_secret=[应用Secret Key]"
    #水利法规问答模型，（微淘之后）
    # API_KEY = "7uad3xLbT4tB4vy6i74EILqL"
    # SECRET_KEY = "BxJ9oUQha5drTM5SJISnib0lDfKpATPF"
    #  水利法规知识库
    # API_KEY = "5UMcxwqv1XLBAlq2pSSdyvjq"
    # SECRET_KEY = "PwLDquDp5CcbPW9uMHD4qJPnHGUOPs6D"
    url = f"https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id={API_KEY}&client_secret={SECRET_KEY}"
    payload = json.dumps("")
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }
    response = requests.request("POST", url, headers=headers, data=payload)
    return response.json().get("access_token")

#注意：翻墙的时候无法调用百度  api  接口
# def query_question(text):
#     url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/eb-instant?access_token=" + get_access_token()
#     #s = input()
#     # 注意message必须是奇数条
#     payload = json.dumps({
#         "messages": [
#             {
#                 "role": "user",
#                 "content": text
#             }
#         ],
#     })
#     headers = {
#         'Content-Type': 'application/json'
#     }
#
#     res = requests.request("POST", url, headers=headers, data=payload).json()
#     return res['result']

# def query_question(text):
#     llm = Ollama(model="qwen2.5")
#     res = llm(text)
#     return res
def query_question(text):
    # 定义 API 的 URL
    url = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"

    # 定义请求头
    headers = {
        "Content-Type": "application/json",
        "Authorization": "Bearer ff47bc89-2fa3-4cd2-ae4f-49f11cb38cf0"
    }

    # 定义请求体
    data = {
        "model": "doubao-1-5-pro-32k-250115",
        "messages": [
            {"role": "system", "content": "你是人工智能助手."},
            {"role": "user", "content": text}
        ]
    }

    # 发送 POST 请求
    response = requests.post(url, headers=headers, json=data)

    # 检查响应状态码
    if response.status_code == 200:
        # 提取并返回 assistant 的回复内容
        return response.json()['choices'][0]['message']['content']
    else:
        # 返回错误信息
        return f"Error: {response.status_code}, {response.text}"

def qiuxun2Word(context=None):
    if context is None:
        context = json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
    tpl = DocxTemplate("./data/echartsWord_v4.docx")
    doc_qx_name = context['DSCHNAME']
    ybddgucTab = []
    for ent in context['realDataTabel']:
        ybddgucTab.append({"mTime": ent['mTime'], "z": ent['z'], "w": ent['w'], "qi": ent['qi'], "qo": ent['qo']})
    context['ddTab'] = ybddgucTab
    # pprint.pprint(context['ddTab'])
    # doc_qx_name = str(doc_qx_name).replace("-", "").replace("%", "")
    tmp_path = "media/xld_doc/{}.docx".format(hash(doc_qx_name + str(time.time())) & sys.maxsize)
    base64_img = context['visualEchartsImgUrl']
    imgData = re.sub('^data:image/.+;base64,', '', base64_img)
    imgdata = base64.b64decode(imgData)
    filename = 'data/yvei.png'  # I assume you have a way of picking unique filenames
    with open(filename, 'wb') as f:
        f.write(imgdata)
        f.flush()
    context['img'] = InlineImage(tpl, 'data/yvei.png', width=Mm(160))
    jinja_env = jinja2.Environment(autoescape=True)
    tpl.render(context=context, jinja_env=jinja_env)
    tpl.save(tmp_path)
    # updateToc(source_path=tmp_path, target_path=tmp_path)
    return tmp_path

def yushuiqing(context=None):
    if context is None:
        context = json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
    con = (f"当前雨水情势：{context['DFORECASTTIME']}，小浪底库水位{context['XLDrz']}米，西霞院库水位{context['XXYrz']}米，潼关流量{context['TGq']}m³/s，三门峡流量{context['SMXq']}m³/s，西霞院流量{context['XXYinq']}m³/s，花园口流量{context['HYKq']}m³/s，利津流量{context['LJq']}m³/s。"
           f"\n预报雨水情势：根据共享接入的黄委水文局预报数据，在{context['forecastTime']}—{context['endTime']}，潼关水文站最大流量{context['maxFlow']}（m³/s）{context['maxFlowTime']},水量{context['waterYieldW']}（亿m）。小浪底水库若按当前调令{context['orderNum']}（m³/s）下泄，"
           f"根据防汛调度业务系统内置的水库调度模型计算结果，库水位变化范围是{context['waterLevel1']}（m）（{context['forecastTime']}）—{context['waterLevel2']}（m）（{context['endTime']}）。")
    llm = Ollama(model="qwen2")
    prompt = f"优化下述雨水情内容，保留所有数据，并优化内容描述。\n雨水情内容如下：{con}"
    res = llm(prompt)
    return res


def huanghe_yuqing_generate(context=None):
    default_context = (
        "7月24日，黄河中游大部地区降小雨，其中渭河下游、三花区间部分地区及大汶河大部降中到大雨、"
        "部分站暴雨，最大点雨量伊洛河古城站87毫米。"
    )
    # 获取当前日期和时间
    now = datetime.now()
    # 获取当前的年、月和日
    year = now.year
    month = now.month
    day = now.day
    frag = search_fragpacks("雨情实况")
    if frag:
        default_context = frag
    if isinstance(context, dict):
        huanghe_weather = context.get("huanghe_weather")  # 从 context 中获取值
        sanhuanjian_weather = context.get("sanhuanjian_weather")
        yuliang_yiluo = context.get("yuliang_yiluo")

        information = (
            f"黄河中下游天气情况：{huanghe_weather}，"
            f"渭河下游天气情况：{sanhuanjian_weather}，"
            f"伊洛河降雨量：{yuliang_yiluo}"
        )

        prompt = (
            f"参考描述：{default_context}\n"
            f"请模仿上述描述，根据以下已知信息生成雨水实况，并进行优化，不要生成无关信息提示，"
            f"也请不要使用'优化后的描述：'这样的字样。"
            f"确保不包含任何说明性文字。"
            f"\n已知信息：{information}，没有时间信息则不显示时间信息。"
        )
        res = query_question(prompt)
        # print("res:",res)
        return res
    else:
        # 如果 context 不是字典，则使用默认上下文
        return default_context

def huanghe_hedaoshuiqing_generate(context):
    if isinstance(context, dict):
        default_list = context['hdsq']
        df = pd.DataFrame(default_list)
        res = pd2HtmlCSS() + divHtml(df.to_html(index=False))
        # prompt = (
        #     "请将以下列表绘制成表格，表头为：站名\t当日8时流量(m³/s)\t昨日均流量(m³/s)。"
        #     "表格名称为“黄河主要站点流量表”，居中置顶。"
        #     "请确保每列的内容对齐，使用固定宽度，格式如下："
        #     "|   站名     | 当日8时流量(m³/s) | 昨日均流量(m³/s) |"
        #     "|   XXX     |       XXX       |      XXX       |"
        #     "请使用空格填充以保持列宽一致，确保所有内容在表格中居中对齐。"
        #     "仅生成表格，避免出现'好的...'、以下是...'、'下面为您展示...'或'以上是...'等字样，"
        #     "确保不包含任何说明性文字。"
        #     f"\n列表内容如下：{default_list}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
        #     res = df.to_html(index=False)
    else:
        res = query_question(context)
    return res

def huanghe_hedaoshuiqing_generate_dfjson(context):
    if isinstance(context, dict):
        default_list = context['hdsq']
        df = pd.DataFrame(default_list)
        res = df.to_json(orient='records')
        # prompt = (
        #     "请将以下列表绘制成表格，表头为：站名\t当日8时流量(m³/s)\t昨日均流量(m³/s)。"
        #     "表格名称为“黄河主要站点流量表”，居中置顶。"
        #     "请确保每列的内容对齐，使用固定宽度，格式如下："
        #     "|   站名     | 当日8时流量(m³/s) | 昨日均流量(m³/s) |"
        #     "|   XXX     |       XXX       |      XXX       |"
        #     "请使用空格填充以保持列宽一致，确保所有内容在表格中居中对齐。"
        #     "仅生成表格，避免出现'好的...'、以下是...'、'下面为您展示...'或'以上是...'等字样，"
        #     "确保不包含任何说明性文字。"
        #     f"\n列表内容如下：{default_list}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     res = df.to_json(orient='records')
    else:
        res = query_question(context)
    return res


def generate_line_shuikushuiqing_table(data):
    header = "| 水库名称 | 水位（m） | 蓄量（亿m³） | 汛限水位（m） | 相应蓄量（亿m³） | 超蓄水量（亿m³） | 剩余防洪库容（亿m³） |"
    rows = []
    for item in data:
        row = (f"| {item['水库名称']} | {item['水位（m）']} | "
               f"{item['蓄量（亿m³）']} | {item['汛限水位（m）']} | "
               f"{item['相应蓄量（亿m³）']} | {item['超蓄水量（亿m³）']} | "
               f"{item['剩余防洪库容（亿m³）']} |")
        rows.append(row)
    return "\n".join([header] + rows)

def huanghe_shuikushuiqing_generate_dfjson(context):
    if isinstance(context, dict):
        default_list = context['sksq']
        df= pd.DataFrame(default_list)
        res = df.to_json(orient='records')
        # prompt = (
        #     "请将以下水库数据转换为直观的表格形式，表头包括：水库名称、 水位（m）、蓄量（亿m³）、汛限水位（m）、相应蓄量（亿m³）、超蓄水量（亿m³）、剩余防洪库容（亿m³）。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "| 水库名称 | 水位（m） | 蓄量（亿m³） | 汛限水位（m） | 相应蓄量（亿m³） | 超蓄水量（亿m³） | 剩余防洪库容（亿m³） |"
        #     "| XXX    |    XXX   |    XXX     |    XXX     |      XXX      |       XXX     |        XXX       |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{default_list}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     res = df.to_json(orient='records')
    else:
        res = query_question(context)
    return res

def huanghe_shuikushuiqing_generate(context):
    if isinstance(context, dict):
        default_list = context['sksq']
        df = pd.DataFrame(default_list)
        res = pd2HtmlCSS() + divHtml(df.to_html(index=False))
        # prompt = (
        #     "请将以下水库数据转换为直观的表格形式，表头包括：水库名称、 水位（m）、蓄量（亿m³）、汛限水位（m）、相应蓄量（亿m³）、超蓄水量（亿m³）、剩余防洪库容（亿m³）。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "| 水库名称 | 水位（m） | 蓄量（亿m³） | 汛限水位（m） | 相应蓄量（亿m³） | 超蓄水量（亿m³） | 剩余防洪库容（亿m³） |"
        #     "| XXX    |    XXX   |    XXX     |    XXX     |      XXX      |       XXX     |        XXX       |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{default_list}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
        #     res = df.to_html(index=False)
    else:
        res = query_question(context)
    return res

def huanghe_gongqing_generate(context):
    default_context = (
        "调水调沙以来，黄河下游累计有39处工程123道坝出险214次，抢险用石5.08万方，耗资1723.27万元。"
        "其中：河南河段累计有30处工程107道坝出险197次，抢险用石4.09万方，耗资1354.85万元；山东段累计有9处工程17道坝出险18次，抢险用石0.99立方米，耗资377.81万元。"
    )

    if isinstance(context, dict):
        huanghexiayou_xianqing = context["huanghexiayou_xianqing"]  # 后续接口获取
        henan_heduan = context["henan_heduan"]  # 后续接口获取
        shangdong_heduan = context["shangdong_heduan"]

        information = (
            f"黄河下游累计{huanghexiayou_xianqing}，"
            f"河南段累计{henan_heduan}；"
            f"山东段累计{shangdong_heduan}。"
        )
        #print("information:",information)
        # data  = text_table(information)
        # df = pd.DataFrame(data)
        # res = df.to_json(orient="records")
        prompt = (
            f"参考描述：{default_context}\n"
            f"请模仿上述描述，根据以下已知信息生成工程险情实况，不要生成总结或无关信息，"
            f"\n已知信息：{information}"
        )
        res = query_question(prompt)
    else:
        res = default_context
    return res

def huanghe_gongqing_generate_html(context):
    default_context = (
        "调水调沙以来，黄河下游累计有39处工程123道坝出险214次，抢险用石5.08万方，耗资1723.27万元。"
        "其中：河南河段累计有30处工程107道坝出险197次，抢险用石4.09万方，耗资1354.85万元；山东段累计有9处工程17道坝出险18次，抢险用石0.99立方米，耗资377.81万元。"
    )

    if isinstance(context, dict):
        huanghexiayou_xianqing = context["huanghexiayou_xianqing"]  # 后续接口获取
        henan_heduan = context["henan_heduan"]  # 后续接口获取
        shangdong_heduan = context["shangdong_heduan"]

        information = (
            f"黄河下游累计有{huanghexiayou_xianqing}，"
            f"河南段累计有{henan_heduan}，"
            f"山东段累计有{ shangdong_heduan}"
        )
        # data  = text_table(information)
        # df = pd.DataFrame(data)
        # res = pd2HtmlCSS() + divHtml(df.to_html(index=False))
        prompt = (
            f"参考描述：{default_context}\n"
            f"请模仿上述描述，根据以下已知信息生成工程险情实况，不要生成总结或无关信息，"
            f"\n已知信息：{information}"
        )
        res = query_question(prompt)
    else:
        res = default_context
    return res

def realtime_yushuigongqing(context):

    yuqing = huanghe_yuqing_generate(context)
    hedaoshuiqing = huanghe_hedaoshuiqing_generate(context)
    shuikushuiqing= huanghe_shuikushuiqing_generate(context)
    gongqing = huanghe_gongqing_generate(context)
    yushuigongqing = (f"1   实时雨水工情\n"
                      f"1.1 雨晴实况   \n"
                      f"\t{yuqing}\n"
                      f"1.2  水情实况   \n"
                      f"1.2.1 河道水情  \n"
                      f"\t{hedaoshuiqing}\n"
                      f"1.2.2 水库水情  \n"
                      f"\t{shuikushuiqing}\n"
                      f"1.3  工情险情实况\n"
                      f"\t{gongqing}")
    return yushuigongqing
    # yushuigong_template=


def huanghe_jiangyu13_forecast(context):
    default_context=("7月23日，兰州以上、山陕北部、泾渭河大部及兰托区间部分地区降小雨，局部中到大雨，个别站暴雨，"
                     "最大点雨量泾河半个城站73毫米；大汶河普降中雨，局地大到暴雨，最大点雨量船厂村站84毫米。\n"
                     "7月24日，兰州以上、沁河、大汶河大部，兰托区间、山陕区间局部地区有小雨；泾渭洛河大部有小雨，局部中雨。\n"
                     "7月25日，兰州以上大部地区有小雨，局部中雨；兰托区间、山陕北部局部小到中雨；泾渭河、伊洛河大部地区有小雨。\n"
                     "7月26日，黄河上中游大部地区有小雨，局部中到大雨；金堤河大部小雨.\n")
    if isinstance(context, dict):

        now = datetime.now()
        # 获取当前的年、月和日
        month = now.month
        day = now.day
        current_jianyu = context["current_jiangyu"]#后续接口获取

        first_jiangyu = context["first_jiangyu"]#"兰州以上、沁河、大汶河大部，兰托区间、山陕区间局部地区有小雨；泾渭洛河大部有小雨，局部中雨。"#后续接口获取
        second_jiangyu = context["second_jiangyu"]#"兰州以上大部地区有小雨，局部中雨；兰托区间、山陕北部局部小到中雨；泾渭河、伊洛河大部地区有小雨。"#后续接口获取
        third_jiangyu = context["third_jiangyu"]#兰州以上大部地区有小雨，局部中雨；兰托区间、山陕北部局部小到中雨；"#后续接口获取
        #print("第三日降雨：", third_jiangyu)

        prompt = (f"参考降雨模板：{default_context}\n"
                  f"请描述当前日期及未来3天的降雨预报，确保使用实际结果替换对应内容，不要添加任何引导性语言或总结性语言。\n"
                  f"当日降雨情况：{current_jianyu}\n"
                  f"未来第1天降雨预报情况：{first_jiangyu}\n"
                  f"未来第2天降雨预报情况：{second_jiangyu}\n"
                  f"未来第3天降雨预报情况：{third_jiangyu}\n")
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_fenqu_jiangyu_forecast(context=None):
    if isinstance(context, dict):
        default_list = context["hhlyjyyb"],
        prompt = (
            "请将以下水库数据转换为直观的表格形式，表头包括：区域、 0-24小时、24-48小时 、 48-72小时。"
            "请确保以下内容对齐并居中显示，格式如下："
            "| 区域    |   0-24小时  | 24-48小时   |   48-72小时  |"
            "| XXX    |    XXX     |    XXX      |    XXX      |"
            "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
            "请确保返回的只有Markdown形式的语法"
            f"\n数据如下：{default_list}"
        )
        #print("prompt:",prompt)
        res = query_question(prompt)
        markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # 使用 StringIO 读取 Markdown 表格
        if markdown_table:  # 确保有匹配的表格
            # 使用 StringIO 读取 Markdown 表格
            data = StringIO(''.join(markdown_table))
            # 读取为 DataFrame
            df = pd.read_csv(data, sep='|', skipinitialspace=True)
            # 清洗数据，移除空列
            df = df.dropna(axis=1, how='all')
            # 使用 tabulate 输出带边框的表格
            final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
            res = final_output
            res = pd2HtmlCSS() + df.to_html(index=False)
    else:
        res = query_question(context)
    return res


def huanghe_fenqu_jiangyu_forecast_json(context=None):
    if isinstance(context, dict):
        default_list = context['hhlyjyyb']
        print("default_list:",default_list)
        df = pd.DataFrame(default_list, columns=["区域", "0-24小时", "24-48小时", "48-72小时"])
        res = df.to_json(orient='records', force_ascii=False)  # 转换为 JSON 格式
def huanghe_fenqu_jiangyu_forecast_dfjson(context=None):
    if isinstance(context, dict):
        default_list = context["hhlyjyyb"]
        df =pd.DataFrame(default_list)
        res = df.to_json(orient='records')
        # prompt = (
        #     "请将以下水库数据转换为直观的表格形式，表头包括：区域、 0-24小时、24-48小时 、 48-72小时。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "| 区域    |   0-24小时  | 24-48小时   |   48-72小时  |"
        #     "| XXX    |    XXX     |    XXX      |    XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{default_list}"
        # )
        # #print("prompt:",prompt)
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     return df.to_json(orient='records')
    else:
        res = query_question(context)
    return res

def  huanghe_jiangyu47_forecast(context):
    default_context=("7月27日，除兰托区间外，黄河流域大部地区有小雨，其中兰州以上局部中雨。\n"
                     "7月28日，黄河流域大部地区有小雨，其中泾渭河上游部分中雨。\n"
                     "7月29日，黄河流域大部地区有小雨，其中山陕区间、伊洛河部分，黄河下游大部有中雨。\n"
                     "7月30日，黄河中下游大部地区小雨，其中三花间部分中雨，黄河下游大部大雨、局部暴雨。\n")
    if isinstance(context, dict):
        now = datetime.now()
        # 获取当前的年、月和日
        month = now.month
        day = now.day
        fourth_jiangyu = context["fourth_jiangyu"]#"除兰托区间外，黄河流域大部地区有小雨，其中兰州以上局部中雨。"#后续接口获取
        #print("第四天降雨预报：", fourth_jiangyu)
        fifth_jiangyu =  context["fifth_jiangyu"]#"黄河流域大部地区有小雨，其中泾渭河上游部分中雨。"#后续接口获取
        sixth_jiangyu = context["sixth_jiangyu"]#"黄河流域大部地区有小雨，其中山陕区间、伊洛河部分，黄河下游大部有中雨。"#后续接口获取
        seventh_jiangyu =  context["seventh_jiangyu"]# "黄河中下游大部地区小雨，其中三花间部分中雨，黄河下游大部大雨。"#后续接口获取
        prompt = (f"参考降雨模板：{default_context}\n"
                  f"请描述当前日期及未来4-7天的降雨预报，确保使用实际结果替换对应内容，不要添加任何引导性语言或总结性语言。\n"
                  f"未来第4天降雨预报情况：：{fourth_jiangyu}\n"
                  f"未来第5天降雨预报情况：：{fifth_jiangyu}\n"
                  f"未来第6天降雨预报情况：：{sixth_jiangyu}\n"
                  f"未来第7天降雨预报情况：：{seventh_jiangyu}\n")

        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_flood_forecast(context):
    if isinstance(context, dict):
        default_list = context["hhfloodforecast"]
        #print("预报水情：",default_list)
        df = pd.DataFrame(default_list)
        res = pd2HtmlCSS() + df.to_html(index=False, justify="center")
        # prompt = (
        #     "请将以下水库数据转换为直观的表格形式，表头包括：日 期、 日均流量（m³/s）、日均流量（m³/s）、日均流量（m³/s）、日均流量（m³/s）。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "| 日 期    |   日均流量（m³/s）  |   日均流量（m³/s）   |   日均流量（m³/s）  |    日均流量（m³/s）  |"
        #     "| XXX     |        XXX        |         XXX       |          XXX      |          XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{default_list}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res
def huanghe_flood_forecast_json(context):
    if isinstance(context, dict):
        default_list = context["hhfloodforecast"]
        #print("预报水情：",default_list)
        df = pd.DataFrame(default_list)
        res = df.to_json(orient='records', force_ascii=False)
    else:
        res = query_question(context)
    return res

def huanghe_yubao(context):

    jiangyu13 = huanghe_jiangyu13_forecast(context)
    jiangyu_table = huanghe_fenqu_jiangyu_forecast(context)
    jiangyu47 = huanghe_jiangyu47_forecast(context)
    flood_forecast = huanghe_flood_forecast(context)
    yubao = (f"1） 降雨预报   \n"
            f"\t{jiangyu13}\n"
            f"2） 分区面平均雨量预报   \n"
            f"黄河流域分区面平均雨量预报（单位：mm）  \n"
            f"\t{jiangyu_table}\n"
            f"3）未来4—7天降水预报  \n"
            f"\t{jiangyu47}\n"
            f"4）洪水预报  \n"
            f"黄河上游未来七天径流预报\n"
            f"\t{flood_forecast}\n"
            )

    return yubao


def huanghe_diaodu_plan(context):
    from yaapp import smx_sk, xld_sk, lh_sk, gx_sk, hkc_sk
    shuiku_shuiwei = {}
    shuiwenzhan_liuliang = {}

    default_context=("按照《黄河洪水调度方案》确定的水库群分级调度规则，建议中游五库调度原则如下"
                     f"三门峡水库：{smx_sk(shuiku_shuiwei, shuiwenzhan_liuliang)},"
                     f"小浪底库：{xld_sk(shuiku_shuiwei, shuiwenzhan_liuliang)},"
                     f"陆浑：{lh_sk(shuiku_shuiwei, shuiwenzhan_liuliang)}"
                     f"故县：{gx_sk(shuiku_shuiwei, shuiwenzhan_liuliang)}"
                     f"河口村库：{hkc_sk(shuiku_shuiwei, shuiwenzhan_liuliang)}"
                     "(TODO)根据建议的调度原则和水库群调度模型计算结果，提出以下调度建议： "
                     "1. （1）三门峡水库按不超305m进出库平衡运用。 "
                     "（2）小浪底水库继续按500m³/s控泄，库水位维持在220m附近。 "
                     "（3）陆浑水库18时起按300m³/s下泄."
                     "（4）故县水库18时起按300m³/s下泄。"
                     "（5）河口村水库18时起按300m³/s下泄。")
    if isinstance(context, dict):
        diaoduyuanze= context["hhddyz"]
        diaodujianyi=context['hhddjy']
        # ddfa_path = context['ddfa_path']
        # print(ddfa_path)
        # tab_html = ddfa_html_table(ddfa_path)
        # tab_html = pd.read_excel(ddfa_path).to_html(index=False)  
        # 不在展示调度方案单的内容
        # prompt = (f"参照调度模板{default_context}\n"
        #           f"根据下述调度原则及调度建议，生成调度方案，不要生成总结或无关信息,以纯文本形式展示,不要显示markdown格式。\n"
        #           f"调度原则为:{diaoduyuanze}\n"
        #           f"调度建议为：{diaodujianyi}\n"
        #           f"按照普通文本输出，不要输出markdown格式")
        # res = query_question(prompt)
        # res = res + "\n"
        res =   (f"调度原则为:{diaoduyuanze}\n"
                f"调度建议为：{diaodujianyi}\n")
    else:
        res = query_question(context)

    return res

def huanghe_diaodu_plan_dfjson(context):
    if isinstance(context, dict):
        ddfa_path = context['ddfa_path']
        # tab_html = ddfa_html_table(ddfa_path)
        return pd.read_excel(ddfa_path).to_json(orient='records')
    else:
        res = query_question(context)
    return res


def huanghe_diaodu_plan_ctx(context):
    default_context=("按照《黄河洪水调度方案》确定的水库群分级调度规则，建议中游五库调度原则如下"
                     "三门峡水库：视潼关站来水来沙情况，原则上按敞泄运用"
                     "小浪底库：原则上按控制花园口站4500m³/s方式运用。若洪水主要来源于三门峡以上，视来水来沙及水库淤积情况，适时按进出库平衡方式运用。控制水库最高运 小"
                     "用水位不超过254m。西霞院水库配合小浪底水库泄洪排沙"
                     "陆浑：按进出库平衡或敞泄运用"
                     "故县：按进出库平衡或敞泄运用"
                     "河口村库：按进出库平衡方式运用。张峰水库适时配合河口村水库拦洪运用"
                     "(TODO)根据建议的调度原则和水库群调度模型计算结果，提出以下调度建议： "
                     "1. （1）三门峡水库按不超305m进出库平衡运用。 "
                     "（2）小浪底水库继续按500m³/s控泄，库水位维持在220m附近。 "
                     "（3）陆浑水库18时起按300m³/s下泄."
                     "（4）故县水库18时起按300m³/s下泄。"
                     "（5）河口村水库18时起按300m³/s下泄。")
    if isinstance(context, dict):
        diaoduyuanze= context["hhddyz"]
        diaodujianyi=context['hhddjy']
        # tab_html = ddfa_html_table(ddfa_path)
        prompt = (f"参照调度模板{default_context}\n"
                  f"根据下述调度原则及调度建议，生成调度方案，不要生成总结或无关信息,以纯文本形式展示。。\n"
                  f"调度原则为:{diaoduyuanze}\n"
                  f"调度建议为：{diaodujianyi}\n")
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_diaodu_plan_yuanze_ctx(context):
    if isinstance(context, dict):
        diaoduyuanze= context["hhddyz"]
        data = extract_shuiku_data(diaoduyuanze)
        df = pd.DataFrame(data)
        res = df.to_json(orient="records")
        #res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_diaodu_plan_yuanze_html(context):
    if isinstance(context, dict):
        diaoduyuanze= context["hhddyz"]
        data = extract_shuiku_data(diaoduyuanze)
        df = pd.DataFrame(data)
        res = pd2HtmlCSS() + df.to_html(index=False, justify="center")
        #res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_diaodu_plan_jianyi_ctx(context):
    if isinstance(context, dict):
        diaodujianyi = context["hhddjy"]
        data = extract_shuiku_data_jianyi(diaodujianyi)
        df = pd.DataFrame(data)
        res = df.to_json(orient="records")
        #res = query_question(prompt)
    else:
        data = extract_shuiku_data_jianyi(context)
        df = pd.DataFrame(data)
        res = df.to_json(orient="records")
        #res = query_question(context)
    return res

def huanghe_diaodu_plan_jianyi_html(context):
    if isinstance(context, dict):
        diaodujianyi= context["hhddjy"]
        data = extract_shuiku_data_jianyi(diaodujianyi)
        df = pd.DataFrame(data)
        res = pd2HtmlCSS() + df.to_html(index=False, justify="center")
        #res = query_question(prompt)
    else:
        data = extract_shuiku_data_jianyi(context)
        df = pd.DataFrame(data)
        res = pd2HtmlCSS() + df.to_html(index=False, justify="center")
        #res = query_question(context)
    return res

def generate_markdown_table(data):
    # 获取表头
    headers = [
        "日期", "潼关预报流量", "三门峡入库流量", "三门峡出库流量",
        "三门峡水位", "三门峡蓄水量", "小浪底入库流量", "小浪底出库流量",
        "小浪底水位", "小浪底蓄水量", "陆浑入库流量", "陆浑出库流量",
        "陆浑水位", "陆浑蓄水量", "故县库流量", "故县出库流量",
        "故县水位", "故县蓄水量", "河口村入库流量", "河口村出库流量",
        "河口村水位", "河口村蓄水量", "花园口流量"
    ]

    # 创建表头
    markdown = "| " + " | ".join(headers) + " |\n"
    markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"

    # 填充数据
    for entry in data:
        row = "| " + " | ".join(str(entry[header]) for header in headers) + " |\n"
        markdown += row

    return markdown


def generate_aligned_markdown_table(data):
    # 定义表头
    headers = [
        "日期", "潼关预报流量", "三门峡入库流量", "三门峡出库流量",
        "三门峡水位", "三门峡蓄水量", "小浪底入库流量", "小浪底出库流量",
        "小浪底水位", "小浪底蓄水量", "陆浑入库流量", "陆浑出库流量",
        "陆浑水位", "陆浑蓄水量", "故县库流量", "故县出库流量",
        "故县水位", "故县蓄水量", "河口村入库流量", "河口村出库流量",
        "河口村水位", "河口村蓄水量", "花园口流量"
    ]

    # 固定列宽
    fixed_width = 20  # 修改为20以提供更多空间

    # 创建表头
    header_row = "| " + " | ".join(header.center(fixed_width) for header in headers) + " |"

    # 创建分隔行
    separator_row = "| " + " | ".join(['-' * fixed_width for _ in headers]) + " |"

    # 填充数据
    rows = []
    for entry in data:
        row = "| " + " | ".join(
            str(entry[header]).center(fixed_width) for header in headers
        ) + " |"
        rows.append(row)

    # 组合所有部分
    markdown = "\n".join([header_row, separator_row] + rows)
    return markdown
def huanghe_diaodu_table(context):
    if "黄河中下游" in context:
        default_list = [
            {
                "日期": "2021-10-27 08:00:00","潼关预报流量": 12800,"三门峡入库流量": 1890,"三门峡出库流量": 1460,"三门峡水位": 317.69,"三门峡蓄水量": 5.48,
                "小浪底入库流量": 1530,"小浪底出库流量": 1400,"小浪底水位": 269.31,"小浪底蓄水量": 80.79,"陆浑入库流量": 50,"陆浑出库流量": 50,"陆浑水位": 319.35,"陆浑蓄水量": 6.57,
                "故县库流量": 93,"故县出库流量": 93,"故县水位": 534.8,"故县蓄水量": 6.27,
                "河口村入库流量": 95,"河口村出库流量": 20,"河口村水位": 272.97,"河口村蓄水量": 2.38,"花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 10:00:00","潼关预报流量": 1800,"三门峡入库流量": 1910,"三门峡出库流量": 1460,"三门峡水位": 317.72,"三门峡蓄水量": 5.51,"小浪底入库流量": 1530,"小浪底出库流量": 1400,"小浪底水位": 269.31,"小浪底蓄水量": 80.8,
                "陆浑入库流量": 50,"陆浑出库流量": 50,"陆浑水位": 319.35,"陆浑蓄水量": 6.57,"故县库流量": 92,"故县出库流量": 92,"故县水位": 534.8,"故县蓄水量": 6.27,"河口村入库流量": 95,
                "河口村出库流量": 20,"河口村水位": 273.06,"河口村蓄水量": 2.38,"花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 08:00:00", "潼关预报流量": 12800, "三门峡入库流量": 1890, "三门峡出库流量": 1460,
                "三门峡水位": 317.69, "三门峡蓄水量": 5.48,
                "小浪底入库流量": 1530, "小浪底出库流量": 1400, "小浪底水位": 269.31, "小浪底蓄水量": 80.79,
                "陆浑入库流量": 50, "陆浑出库流量": 50, "陆浑水位": 319.35, "陆浑蓄水量": 6.57,
                "故县库流量": 93, "故县出库流量": 93, "故县水位": 534.8, "故县蓄水量": 6.27,
                "河口村入库流量": 95, "河口村出库流量": 20, "河口村水位": 272.97, "河口村蓄水量": 2.38,
                "花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 10:00:00", "潼关预报流量": 1800, "三门峡入库流量": 1910, "三门峡出库流量": 1460,
                "三门峡水位": 317.72, "三门峡蓄水量": 5.51, "小浪底入库流量": 1530, "小浪底出库流量": 1400,
                "小浪底水位": 269.31, "小浪底蓄水量": 80.8,
                "陆浑入库流量": 50, "陆浑出库流量": 50, "陆浑水位": 319.35, "陆浑蓄水量": 6.57, "故县库流量": 92,
                "故县出库流量": 92, "故县水位": 534.8, "故县蓄水量": 6.27, "河口村入库流量": 95,
                "河口村出库流量": 20, "河口村水位": 273.06, "河口村蓄水量": 2.38, "花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 08:00:00", "潼关预报流量": 12800, "三门峡入库流量": 1890, "三门峡出库流量": 1460,
                "三门峡水位": 317.69, "三门峡蓄水量": 5.48,
                "小浪底入库流量": 1530, "小浪底出库流量": 1400, "小浪底水位": 269.31, "小浪底蓄水量": 80.79,
                "陆浑入库流量": 50, "陆浑出库流量": 50, "陆浑水位": 319.35, "陆浑蓄水量": 6.57,
                "故县库流量": 93, "故县出库流量": 93, "故县水位": 534.8, "故县蓄水量": 6.27,
                "河口村入库流量": 95, "河口村出库流量": 20, "河口村水位": 272.97, "河口村蓄水量": 2.38,
                "花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 10:00:00", "潼关预报流量": 1800, "三门峡入库流量": 1910, "三门峡出库流量": 1460,
                "三门峡水位": 317.72, "三门峡蓄水量": 5.51, "小浪底入库流量": 1530, "小浪底出库流量": 1400,
                "小浪底水位": 269.31, "小浪底蓄水量": 80.8,
                "陆浑入库流量": 50, "陆浑出库流量": 50, "陆浑水位": 319.35, "陆浑蓄水量": 6.57, "故县库流量": 92,
                "故县出库流量": 92, "故县水位": 534.8, "故县蓄水量": 6.27, "河口村入库流量": 95,
                "河口村出库流量": 20, "河口村水位": 273.06, "河口村蓄水量": 2.38, "花园口流量": 1720,
            },{
                "日期": "2021-10-27 08:00:00","潼关预报流量": 12800,"三门峡入库流量": 1890,"三门峡出库流量": 1460,"三门峡水位": 317.69,"三门峡蓄水量": 5.48,
                "小浪底入库流量": 1530,"小浪底出库流量": 1400,"小浪底水位": 269.31,"小浪底蓄水量": 80.79,"陆浑入库流量": 50,"陆浑出库流量": 50,"陆浑水位": 319.35,"陆浑蓄水量": 6.57,
                "故县库流量": 93,"故县出库流量": 93,"故县水位": 534.8,"故县蓄水量": 6.27,
                "河口村入库流量": 95,"河口村出库流量": 20,"河口村水位": 272.97,"河口村蓄水量": 2.38,"花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 10:00:00","潼关预报流量": 1800,"三门峡入库流量": 1910,"三门峡出库流量": 1460,"三门峡水位": 317.72,"三门峡蓄水量": 5.51,"小浪底入库流量": 1530,"小浪底出库流量": 1400,"小浪底水位": 269.31,"小浪底蓄水量": 80.8,
                "陆浑入库流量": 50,"陆浑出库流量": 50,"陆浑水位": 319.35,"陆浑蓄水量": 6.57,"故县库流量": 92,"故县出库流量": 92,"故县水位": 534.8,"故县蓄水量": 6.27,"河口村入库流量": 95,
                "河口村出库流量": 20,"河口村水位": 273.06,"河口村蓄水量": 2.38,"花园口流量": 1720,
            },{
                "日期": "2021-10-27 08:00:00","潼关预报流量": 12800,"三门峡入库流量": 1890,"三门峡出库流量": 1460,"三门峡水位": 317.69,"三门峡蓄水量": 5.48,
                "小浪底入库流量": 1530,"小浪底出库流量": 1400,"小浪底水位": 269.31,"小浪底蓄水量": 80.79,"陆浑入库流量": 50,"陆浑出库流量": 50,"陆浑水位": 319.35,"陆浑蓄水量": 6.57,
                "故县库流量": 93,"故县出库流量": 93,"故县水位": 534.8,"故县蓄水量": 6.27,
                "河口村入库流量": 95,"河口村出库流量": 20,"河口村水位": 272.97,"河口村蓄水量": 2.38,"花园口流量": 1720,
            },
            {
                "日期": "2021-10-27 10:00:00","潼关预报流量": 1800,"三门峡入库流量": 1910,"三门峡出库流量": 1460,"三门峡水位": 317.72,"三门峡蓄水量": 5.51,"小浪底入库流量": 1530,"小浪底出库流量": 1400,"小浪底水位": 269.31,"小浪底蓄水量": 80.8,
                "陆浑入库流量": 50,"陆浑出库流量": 50,"陆浑水位": 319.35,"陆浑蓄水量": 6.57,"故县库流量": 92,"故县出库流量": 92,"故县水位": 534.8,"故县蓄水量": 6.27,"河口村入库流量": 95,
                "河口村出库流量": 20,"河口村水位": 273.06,"河口村蓄水量": 2.38,"花园口流量": 1720,
            },
            # 其他数据...
        ]
        prompt = (
            "请将以下水库数据转换为直观的表格形式，表头包括：日 期、 潼关预报流量、三门峡入库流量、三门峡出库流量、三门峡水位、三门峡蓄水量、小浪底入库流量、小浪底出库流量、小浪底水位、小浪底蓄水量、陆浑入库流量、陆浑出库流量、陆浑水位、陆浑蓄水量、故县库流量、故县出库流量、故县水位、故县蓄水量、河口村入库流量、河口村出库流量、河口村水位、河口村蓄水量、花园口流量。"
            "请确保以下内容对齐并居中显示，格式如下："
            "|日 期|三门峡入库流量|三门峡出库流量|三门峡水位|三门峡蓄水量|小浪底入库流量|小浪底出库流量|小浪底水位|小浪底蓄水量|陆浑入库流量|陆浑出库流量|陆浑水位|陆浑蓄水量|故县入库流量|故县出库流量|故县水位|故县蓄水量|河口村入库流量|河口村出库流量|河口村水位|河口村蓄水量|花园口流量|"
            "| XXX|    XXX    |   XXX     |  XXX   |   XXX    |   XXX     |   XXX     |   XXX  |   XXX   |    XXX   |   XXX   |  XXX  |  XXX   |   XXX   |    XXX   |  XXX |  XXX   |  XXX      |  XXX      |  XXX   |  XXX    |  XXX   |"
            "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
            "请确保返回的只有Markdown形式的语法"
            f"\n数据如下：{default_list}"
        )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
        #res = generate_markdown_table(default_list)
        res = generate_aligned_markdown_table(default_list)
    else:
        res = query_question(context)
    return res


def huanghe_diaodu(context=None):
    diaodu_plan=huanghe_diaodu_plan(context)
    diaodu= (f"3 调度方案\n"
             f"{diaodu_plan}\n")
    return diaodu

def huanghe_shuiku_diaodu_result(context):
    default_context = ("预计小浪底水库将于2021年10月30日12时达到最高水位269.99m，超汛限水位34.99m；\n"
                       "预计三门峡水库将于2021年10月28日02时达到最高水位317.96m，低于人员紧急转移水位0.04m；\n"
                       "预计陆浑水库将于2021年10月27日08时达到最高水位319.35m，低于征地水位0.15m；\n"
                       "预计故县水库将于2021年10月27日08时达到最高水位534.8m，超征地水位1.16m；\n"
                       "预计河口村水库将于2021年11月02日20时达到最高水位272.97m，超汛限水位34.97m；\n")
    if isinstance(context, dict):
        xiaolangdi_result = context["xldddresult"]
        sanmenxia_result = context["smxddresult"]
        luhun_result = context["lhddresult"]
        guxian_result = context["gxddresult"]
        hekoucun_result = context["hkcddresult"]

        prompt = (f"参照调度结果模板：\n{default_context}\n"
                  f"请根据以下调度结果生成描述，确保使用实际结果替换对应内容，不要生成总结或无关信息。\n"
                  f"小浪底水库调度结果为: {xiaolangdi_result}\n"
                  f"三门峡水库调度结果为: {sanmenxia_result}\n"
                  f"陆浑水库调度结果为: {luhun_result}\n"
                  f"故县水库调度结果为: {guxian_result}\n"
                  f"河口村水库调度结果为: {hekoucun_result}\n"
                  f"请输出完整的调度结果描述，确保结果与实际相符。"
                  )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_tanqu_yanmo(context):
    default_context = ("根据《山东省黄河滩区运用预案》、《河南省黄河滩区运用预案》，在花园口6000流量下，"
                       "河南省预计进水村庄70个、人口6.18万人，水围村庄273个、人口24.85万人，淹没滩地110.29万亩，淹没耕地70.81万亩，经济损失140.87亿元；"
                       "山东省预计漫滩面积万亩，淹没耕地39.07万亩，93个滩区进水，18个自然村进水，46个自然村被水围困，涉及4.54万人，需转移安置2.46万人，就地或就近安置2.09万人。")
    if isinstance(context, dict):
        henan_yanmo = context["henan_yanmo"]
        shandong_yanmo = context["shandong_yanmo"]

        prompt = (f"参照模板：\n{default_context}\n"
                  f"请根据以下滩区淹没结果生成描述，确保使用实际结果替换对应内容，不要生成总结或无关信息，,以纯文本形式展示。\n"
                  f"河南省滩区淹没结果为: {henan_yanmo}\n"
                  f"山东省滩区淹没结果为: {shandong_yanmo}\n"
                  f"请输出完整的调度结果描述，确保包含背景信息和实际结果。"
                  )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_keneng_danger(context):
    default_context = ("根据《山东省黄河防洪预案》《河南省黄河防洪预案》，在花园口6000流量级下，河南、山东两省工情、险情预测预测如下：河南省：部分工程已接近或超过工程防洪标准，洪水沿主河槽下泄，"
                       "部分低滩区可能漫水，此级洪水极易引起河势上提或下挫，坐湾生险，河道整治工程出险几率增大。；山西省：本级洪水冲刷力较强，局部河段河势可能发生较大变化，如东明县辛店集工程、堡城险工下首，"
                       "鄄城县苏泗庄险工下首，东平县荫柳科控导、长清区娘娘店控导上首，惠民县五甲杨险工上首，垦利区十八户控导、义和险工、卞庄险工上首，河口区清三控导上首；"
                       "险工、控导工程特别是新修工程易发生根石走失及坦石坍塌、墩蛰等险情，如近几年修建的垦利区十八户上首、东营区老于滩、章丘区王家圈等应急防护工程，"
                       "改建的牡丹区刘庄险工、牡丹区贾庄险工、齐河县王庄险工，东明县辛店集工程、王高寨控导、老君堂控导、鄄城县郭集控导、苏泗庄控导、阳谷县陶城铺控导等工程，"
                       "新建的鄄城县桑庄潜坝、平阴县胜利控导等工程，东明县霍寨险工、梁山县路那里险工、惠民县大崔险工、惠民县王平口控导、垦利区卞庄险工、垦利区苇改闸控导工程等脱河多年后又重新靠溜；"
                       "洪峰较大时，部分控导工程可能漫顶。漫滩偎堤堤段可能出现风浪淘刷险情，薄弱堤段可能出现渗水、管涌等险情。洪水漫滩或堤沟河行洪，会切断控导工程抢险和群众撤离道路。")
    if isinstance(context, dict):
        henan_keneng_danger = context["henan_keneng_danger"]
        shandong_keneng_danger = context["shandong_keneng_danger"]

        prompt = (f"参照模板：\n{default_context}\n"
                  f"请根据以下工情、险情生成描述，确保使用实际结果替换对应内容，不要生成总结或无关信息,以纯文本形式展示。\n"
                  f"河南省工情、险情预测结果为: {henan_keneng_danger}\n"
                  f"山东省滩区淹没结果为: {shandong_keneng_danger}\n"
                  f"请输出完整的调度结果描述，确保包含背景信息和实际结果。"
                  )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def huanghe_xiangying_level(context):
    default_context = [{"一级预警":"按照《黄河防汛抗旱应急预案》，启动一级应急响应，响应行动如下：（1）黄河防总总指挥或常务副总指挥坐镇指挥黄河抗洪工作，主持抗洪抢险会商会，研究部署抗洪抢险工作。视情与相关省区进行异地会商。（2）根据会商意见，黄河防总办公室向相关省区防指通报关于启动防汛一级应急响应的命令及黄河汛情，对防汛工作提出要求，并向黄河防总总指挥报告。黄河防总向国家防总、水利部报告有关情况，为国家防总和水利部提供调度参谋意见，请求加强对黄河抗洪抢险指导，动员社会力量支援黄河抗洪抢险救灾。（3）黄河防总办公室各成员单位按照黄委防御大洪水职责分工和机构设置上岗到位，全面开展工作，各职能组充实人员。黄委全体职工全力投入抗洪抢险工作。水情测报组滚动进行洪水预测预报，每日至少制作发布气象水情预报 3 次，每日至少提供12 次干支流重要测站监测信息，情况紧急时根据需要加密测报；综合调度组根据预报滚动计算水利工程调度方案，做好干流及重要支流水库调度和东平湖、北金堤滞洪区运用的分析研判；宣传组适时举行新闻发布会，向社会报道黄河抗洪抢险动态，做好新闻宣传工作。（4）黄河防总根据汛情需要，及时增派司局级领导带队的工作组、专家组赶赴现场，指导抗洪抢险救灾工作。（5）根据各地抗洪抢险需要，黄河防总按程序调度黄委防汛物资、黄河机动抢险队支援抗洪抢险，必要时请求国家防总调动流域内外抢险队、物资支援黄河抗洪抢险。（6）有关省区防汛抗旱指挥机构的主要负责同志主持会商，动员部署防汛工作；按照权限组织调度水工程；根据预案转移安置危险地区群众，组织强化巡堤查险和堤防防守，及时控制险情；增派工作组、专家组赴一线指导防汛工作；受灾地区的各级防汛抗旱指挥机构负责人、成员单位负责人，应按照职责到分管的区域组织指挥防汛工作，或驻点具体帮助重灾区做好防汛工作；可按照预案和程序适时请调人民解放军和武警部队支援黄河抗洪抢险；将工作情况上报省区人民政府及黄河防总。根据汛情，相关县级以上人民政府防汛抗旱指挥部宣布进入紧急防汛期，动员一切社会力量投入黄河抗洪抢险"},
                       {"二级预警":"按照《黄河防汛抗旱应急预案》，启动二级应急响应，响应行动如下：黄河防总总指挥或常务副总指挥坐镇指挥黄河抗洪工作，主持抗洪抢险会商会，研究部署抗洪抢险工作。2小时内到达现场指挥工作。"},
                       {"三级预警":"按照《黄河防汛抗旱应急预案》，启动三级应急响应，响应行动如下：黄河防总总指挥或常务副总指挥坐镇指挥黄河抗洪工作，主持抗洪抢险会商会，研究部署抗洪抢险工作。3小时内到达现场指挥工作。"},
                       {"四级预警": "按照《黄河防汛抗旱应急预案》，启动四级应急响应，响应行动如下：黄河防总总指挥或常务副总指挥坐镇指挥黄河抗洪工作，主持抗洪抢险会商会，研究部署抗洪抢险工作。4小时内到达现场指挥工作。"}]
    if isinstance(context, dict):

        yujing_level="三级预警"
        prompt = (f"各个预警等级对应的响应措施如下：{default_context}。\n"
                  f"已知预警等级为：{yujing_level}。\n"
                  f"请根据一直预警等级生成对应的响应措施，不要生成总结或无关信息。")
        res = query_question(prompt)
        #res = yujingdengji(shuiku_shuiwei: dict, shuiwenzhan_liuliang: dict)
    else:
        res = query_question(context)
    return res

def huanghe_diaodu_result(context):
    shuiku_diaodu_result=huanghe_shuiku_diaodu_result(context)
    tanqu_yanmo_result=huanghe_tanqu_yanmo(context)
    keneng_danger=huanghe_keneng_danger(context)
    xiangying_level=huanghe_xiangying_level(context)
    #xiangying_level =
    diaodu_result=(f"\n"
                   f"1) 水库 \n"
                   f"{shuiku_diaodu_result}\n"
                   f"2) 滩区淹没\n"
                   f"{tanqu_yanmo_result}\n"
                   f"3) 可能出险情况\n"
                   f"{keneng_danger}\n"
                   f"4) 预警及响应等级\n"
                   f"{xiangying_level}"
                   )
    return diaodu_result


def xld_yushui_context(context=None):
    default_context=("\t当前雨水情势：3日 23时，小浪底库水位269.890米，西霞院库水位133.360米，潼关流量2720m3/s，三门峡流量2750m3/s，西霞院流量2750m3/s，花园口流量4850m3/s，利津流量5070m3/s。\n"
                     "\t预报雨水情势：根据共享接入的黄委水文局预报数据，在05日 08时—12日 08时，潼关水文站最大流量10660（m3/s）07日 06时,水量78.06（亿m）。小浪底水库若按当前调令3750.0（m3/s）下泄，根据防汛调度业务系统内置的水库调度模型计算结果，库水位变化范围是269.89（m）（05日 08时）—275（m）（12日 08时）。")
    frag =search_fragpacks("雨情实况")
    if frag:
         default_context = frag
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        prompt = (f"参照模板：\n{default_context}\n"
                  f"请根据以下真实数据，确保使用实际结果替换对应内容，不要生成总结或无关信息。\n"
                  f"当前雨水情势： {xld_data['DFORECASTTIME']}\n"
                  f"小浪底库水位： {xld_data['XLDrz']}米\n"
                  f"西霞院库水位：{xld_data['XXYrz']}米"
                  f"潼关流量：{xld_data['TGq']}m³/s"
                  f"三门峡流量：{xld_data['SMXq']}m³/s"
                  f"西霞院流量：{xld_data['XXYinq']}m³/s"
                  f"花园口流量：{xld_data['HYKq']}m³/s"
                  f"利津流量：{xld_data['LJq']}m³/s"
                  f"预报时间：{xld_data['forecastTime']}—{xld_data['endTime']}"
                  f"潼关水文站最大流量：{xld_data['maxFlow']}（m³/s）"
                  f"潼关水文站最大流量发生时间：{xld_data['maxFlowTime']}"
                  f"潼关水文站水量：{xld_data['waterYieldW']}（亿m）"
                  f"当前调令{xld_data['orderNum']}（m3/s）"
                  f"库水位变化范围:{xld_data['waterLevel1']}（m）（{xld_data['forecastTime']}）—{xld_data['waterLevel2']}（m）（{xld_data['endTime']}）"
                  f"请输出完整的调度结果描述，确保包含背景信息和实际结果。"
                  )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def xld_yushuqing(context=None):
    xld_yushui = xld_yushui_context(context)
    yushui_result = (f"一   雨水情势\n"
                     f"{xld_yushui}"
                     )
    return yushui_result



def engineer_safety_shuikuyj(context=None):
    if isinstance(context, dict):
        xld_data = context   #json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        shuikudata=[]
        for shuikudict in xld_data["shuiKuTable"]:
            shuikudata.append({"水库":shuikudict['STNM'],'监测时间':shuikudict['STIME'],'库水位':shuikudict['SRZ'],'超警类型':shuikudict['SPTYPE'],"超警状态":shuikudict['SPSTATUS']})
        df = pd.DataFrame( shuikudata)
        res = pd2HtmlCSS() + df.to_html(index=False)
        return res

        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：水 库、 监测时间、库水位、超警类型、超警状态。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   水 库    |   监测时间）  |   库水位    |   超警类型  |    超警状态  |"
        #     "|   XXX     |    XXX      |   XXX      |    XXX    |     XXX     |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{shuikudata}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res

def engineer_safety_shuikuyj_json(context=None):
    if isinstance(context, dict):
        xld_data = context   #json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        shuikudata=[]
        for shuikudict in xld_data["shuiKuTable"]:
            shuikudata.append({"水库":shuikudict['STNM'],'监测时间':shuikudict['STIME'],'库水位':shuikudict['SRZ'],'超警类型':shuikudict['SPTYPE'],"超警状态":shuikudict['SPSTATUS']})
        df = pd.DataFrame( shuikudata)
        res = df.to_json(orient='records', force_ascii=False)
        return res
    else:
        res = query_question(context)
    return res

def engineer_safety_shuiwenyj(context=None):
    if isinstance(context, dict):
        xld_data = context
        shuiwendata = []
        for shuiwendict in xld_data['shuiWenTable']:
            shuiwendata.append({"水文站":shuiwendict['STNM'],'监测时间':shuiwendict['STIME'],'流量':shuiwendict['STREALVALUE'],'超警类型':shuiwendict['SALARMNAME'],'超警状态':shuiwendict['SALARMVALUE']})
        df = pd.DataFrame(shuiwendata)
        res = pd2HtmlCSS() + df.to_html(index=False)
        return res
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：水文站、 监测时间、流量、超警类型、超警状态。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   水文站    |   监测时间）  |   流量    |   超警类型  |    超警状态  |"
        #     "|   XXX      |    XXX      |   XXX    |    XXX    |     XXX     |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{shuiwendata}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res

def engineer_safety_shuiwenyj_json(context=None):
    if isinstance(context, dict):
        xld_data = context
        shuiwendata = []
        for shuiwendict in xld_data['shuiWenTable']:
            shuiwendata.append({"水文站":shuiwendict['STNM'],'监测时间':shuiwendict['STIME'],'流量':shuiwendict['STREALVALUE'],'超警类型':shuiwendict['SALARMNAME'],'超警状态':shuiwendict['SALARMVALUE']})
        df = pd.DataFrame(shuiwendata)
        res = df.to_json(orient='records', force_ascii=False)
        return res
    else:
        res = query_question(context)
    return res

def engineer_safety_gongchengjcyj(context=None):
    if isinstance(context, dict):
        xld_data = context
        gongchengjiancedata = []
        for gongchengjiancedict in xld_data['jiankongTable']:
            gongchengjiancedata.append({"测点编码":gongchengjiancedict['sensorCode'],'工程部位':gongchengjiancedict['siteNames'],
                                        '监测时间':gongchengjiancedict['meaTime'],'测值':gongchengjiancedict['value'],
                                        '报警等级':gongchengjiancedict['gradeName'],'报警信息':gongchengjiancedict['gradeInfo']})
        df = pd.DataFrame(gongchengjiancedata)
        res = pd2HtmlCSS() + df.to_html(index=False)
        return res
        #
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：测点编码、 工程部位、监测时间、测值、报警等级、报警信息。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   测点编码    |   工程部位  |   监测时间    |   测值     |    报警等级  |  报警信息|"
        #     "|   XXX       |    XXX     |    XXX      |    XXX    |     XXX     |   XXX  |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{gongchengjiancedata}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res
def engineer_safety_gongchengjcyj_json(context=None):
    if isinstance(context, dict):
        xld_data = context
        gongchengjiancedata = []
        for gongchengjiancedict in xld_data['jiankongTable']:
            gongchengjiancedata.append(
                {"测点编码": gongchengjiancedict['sensorCode'], '工程部位': gongchengjiancedict['siteNames'],
                 '监测时间': gongchengjiancedict['meaTime'], '测值': gongchengjiancedict['value'],
                 '报警等级': gongchengjiancedict['gradeName'], '报警信息': gongchengjiancedict['gradeInfo']})
        df = pd.DataFrame(gongchengjiancedata)
        res = df.to_json(orient='records', force_ascii=False)
        return res
    else:
        res = query_question(context)
    return res
def xld_engineer_safety(context=None):
    xld_shukuyujing = engineer_safety_shuikuyj(context)
    xld_shuwenyujing = engineer_safety_shuiwenyj(context)
    xld_gongchengjcyj = engineer_safety_gongchengjcyj(context)
    yushui_result = (f"二   工程安全研判\n"
                     f"水库预警情况表\n"
                     f"{xld_shukuyujing}\n"
                     f"水文站预警情况表\n"
                     f"{xld_shuwenyujing}\n"
                     f"工程监测指标预警情况表\n"
                     f"{xld_gongchengjcyj}\n"
                     )
    return yushui_result


def shuniuFangAn(context=None):
    if isinstance(context, dict):
        default_context = "按照推荐的预演调度方案，05日 07时—12日 08时，小浪底水库按照（m3/s）3562.5—3937.5下泄，库水位变化范围是(m)269.89（05日 08时）—275（12日 08时）。"
        xld_data = context   #json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        prompt=(f"参照模板：\n{default_context}\n"
            f"请根据以下真实数据,结合参照模板，确保使用实际结果替换模板内容，不要生成总结或无关信息。\n"
            f"调度时间：{xld_data['forecastTime']}—{xld_data['endTime']}"
            f"下泄流量：{xld_data['avgflowll']}—{xld_data['avgflowul']}"
            f"库水位变化范围：{xld_data['waterLevel1']}（{xld_data['forecastTime']}）—{xld_data['waterLevel2']}（{xld_data['endTime']}）"
            f"请只输出调度结果，不要添加任何分析、推测或总结。")
        # con = (f"按照推荐的预演调度方案，{xld_data['forecastTime']}—{xld_data['endTime']}，小浪底水库按照（m3/s）{xld_data['avgflowll']}—{xld_data['avgflowul']}下泄，"
        #        f"库水位变化范围是(m){xld_data['waterLevel1']}（{xld_data['forecastTime']}）—{xld_data['waterLevel2']}（{xld_data['endTime']}）")
        #llm = Ollama(model="qwen2")
        #prompt = f"优化下述内容，保留所有数据，并优化内容描述,直接输出优化后的内容。\n内容如下：{con}"
        # prompt = f"请优化以下内容，保留所有数据，并使描述更加流畅，不要生成总结或无关信息。\n{con}"
        res = query_question(prompt)
    else:
        res = query_question(context)
    #res = llm(prompt)
    return res

def llava_call(data):
    url = "http://localhost:11434/api/generate"
    response = requests.post(url, json=data)
    return response.json()
    # 打印响应

def xldshuiddProcessAnalysys(context=None):
    if context is None:
        context = json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
    llm = Ollama(model="llava")
    prompt = f"请从下图中描述小浪底水位变化过程，用汉语回答。"
    base64_img = context['visualEchartsImgUrl']
    imgData = re.sub('^data:image/.+;base64,', '', base64_img)
    data = {"model": "llava", "prompt": prompt,"images":[imgData],'stream':False}
    #res = llm(image_data=imgData,prompt=prompt)
    res = llava_call(data)['response']
    return res

def xldJZStatus(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        JZStatus=[]
        for JZStatusdict in xld_data['genStatusTableData']:
            JZStatus.append({'机组名称':JZStatusdict['gennm'],'机组运行状态':JZStatusdict['gensta']})
        df = pd.DataFrame(JZStatus)
        res = pd2HtmlCSS() + df.to_html(index=False)
        return res
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：机组名称、 机组运行状态。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   机组名称    |   机组运行状态   |"
        #     "|   XXX       |       XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{JZStatus}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output

    else:
        res = query_question(context)
    return res

def xldJZStatus_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        JZStatus=[]
        for JZStatusdict in xld_data['genStatusTableData']:
            JZStatus.append({'机组名称':JZStatusdict['gennm'],'机组运行状态':JZStatusdict['gensta']})
        df = pd.DataFrame(JZStatus)
        res = df.to_json(orient='records', force_ascii=False)
        return res
    else:
        res = query_question(context)
    return res


def xldholeStatus(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        holeStatus=[]
        for holeStatusdict in  xld_data['holeStatusTableData']:
            holeStatus.append({'孔洞名称':holeStatusdict['HONM'],'孔洞运行状态':holeStatusdict['HOSTA']})
        df = pd.DataFrame(holeStatus)
        res = df.to_html(index=False)
        return res
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：机组名称、 机组运行状态。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   机组名称    |   机组运行状态   |"
        #     "|   XXX       |       XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{holeStatus}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res
def xldholeStatus_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        holeStatus=[]
        for holeStatusdict in  xld_data['holeStatusTableData']:
            holeStatus.append({'孔洞名称':holeStatusdict['HONM'],'孔洞运行状态':holeStatusdict['HOSTA']})
        df = pd.DataFrame(holeStatus)
        res = df.to_json(orient='records', force_ascii=False)
        return res
    else:
        res = query_question(context)
    return res

def JZHoleRecommend(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        JZHoleYunYong=(f"按照《黄河水利水电开发集团有限公司2021年防汛抢险应急预案》中的“小浪底水利枢纽泄洪系统孔洞组合运用方案”，结合目前孔洞开启情况，"
                       f"{xld_data['DFORECASTTIME']}，孔洞组合运用方式如下：\n"
                       f"机组、孔洞运用方式推荐表\n")
        JZHoleRec=[]
        for jzhooledict in xld_data['responseData']:
            JZHoleRec.append({'时间':jzhooledict['date'],'开启机组组合':jzhooledict['jizuState'],
                              '时间':jzhooledict['date'],'开启孔洞组合':jzhooledict['kongdongState']})
        df = pd.DataFrame(JZHoleRec)
        res = pd2HtmlCSS() + df.to_html(index=False)
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：时间、 开启机组组合、 时间、 开启孔洞组合。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   时间    |   开启机组组合   |   时间    |   开启孔洞组合   |"
        #     "|   XXX    |       XXX      |   XXX    |      XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{JZHoleRec}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
        res = JZHoleYunYong+res
    else:
        res = query_question(context)
    return res

def JZHoleRecommend_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        JZHoleRec=[]
        for jzhooledict in xld_data['responseData']:
            JZHoleRec.append({'时间':jzhooledict['date'],'开启机组组合':jzhooledict['jizuState'],
                              '时间':jzhooledict['date'],'开启孔洞组合':jzhooledict['kongdongState']})
        df = pd.DataFrame(JZHoleRec)
        res  = df.to_json(orient='records', force_ascii=False)
        return res
    else:
        res = query_question(context)
    return res


def xld_shuniu_apply(context):
    shuniu_apply = shuniuFangAn(context)
    JZStatus  = xldJZStatus(context)
    holeStatus = xldholeStatus(context)
    HoleRecommend = JZHoleRecommend(context)
    shuniu_apply_result = (f"三   枢纽运用方案\n"
                     f"{shuniu_apply}\n"
                     f"预报小浪底水库的调度过程\n"
                     f"小浪底机组状态表\n"
                     f"{JZStatus}\n"
                     f"小浪底孔洞状态表\n"
                     f"{holeStatus}"
                     f"机组、孔洞运用方式推荐\n"
                     f"{HoleRecommend}\n"
                     )
    return shuniu_apply_result


def YingjiResponse(context=None):
    if isinstance(context, dict):
        default_context = "按照《小浪底水利枢纽及西霞院工程 2021年防汛抢险应急预案》、《黄河水利水电开发集团有限公司2021年防汛抢险应急预案》、《水利部小浪底水利枢纽管理中心突发事故应急救援总体预案》、《水利部小浪底水利枢纽管理中心水库调度管理办法》、《水利部小浪底水利枢纽管理中心防汛工作管理办法》、《水利部小浪底水利枢纽管理中心调度（防汛）值班实施细则》等预案要求，在当前潼关流量10661、小浪底水库水位275、小浪底出库流量5340下，启动【二级应急响应】，具体响应措施如下：（1）省防指必要时成立小浪底工程防汛应急抢险前方指挥部，主持召集会议，现场指挥抢险工作；洛阳市、济源市开展本行政区域内的小浪底水利枢纽防汛抢险工作。（2）小浪底管理中心防汛领导小组全体人员在 4 小时内抵达枢纽管理区，配合省防指开展防汛工作。（3）开发公司防汛指挥部全体人员及防汛工作人员在 4 小时内抵达枢纽管理区，按照防汛预案开展巡查监测及通讯后勤保障工作，做好抢险准备。"
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        prompt=(f"参照模板：\n{default_context}\n"
            f"请根据以下真实数据,结合参照模板，确保使用实际结果替换模板内容，不要生成总结或无关信息。\n"
            f"潼关流量：{xld_data['maxFlow']}"
            f"小浪底水库水位：{xld_data['XLDWaterLevel']}"
            f"小浪底出库流量{xld_data['outWaterYield']}"
            f"具体响应措施：{xld_data['resDataSortStr']} "
            f"请只输出调度结果，不要添加任何分析、推测或总结。")
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def OrganizeBaoZhang_leader(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        xldleaderdata = []
        for xldleaderdict in xld_data['XLDheadquarters']:
            xldleaderdata.append({'岗位':xldleaderdict['job_desc'],'责任人':xldleaderdict['name'],'联系方式':xldleaderdict['concat'],'防汛职责':xldleaderdict['duty']})
        df = pd.DataFrame(xldleaderdata)
        res = pd2HtmlCSS() + df.to_html(index=False)
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：岗位、 责任人、 联系方式、 防汛职责。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   岗位    |       责任人     |   联系方式    |      防汛职责   |"
        #     "|   XXX    |        XXX      |     XXX     |       XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{xldleaderdata}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res


def OrganizeBaoZhang_leader_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        xldleaderdata = []
        for xldleaderdict in xld_data['XLDheadquarters']:
            xldleaderdata.append({'岗位':xldleaderdict['job_desc'],'责任人':xldleaderdict['name'],'联系方式':xldleaderdict['concat'],'防汛职责':xldleaderdict['duty']})
        df = pd.DataFrame(xldleaderdata)
        res = df.to_json(orient='records', force_ascii=False)
    else:
        res = query_question(context)
    return res

def OrganizeBaoZhang_zhihuibu(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        xxyYingjidata = []
        for xxyYingjidict in xld_data['XLDheadquarters']:
            xxyYingjidata.append({'岗位':xxyYingjidict['job_desc'],'责任人':xxyYingjidict['name'],'联系方式':xxyYingjidict['concat'],'防汛职责':xxyYingjidict['duty']})
        df = pd.DataFrame(xxyYingjidata)
        res = pd2HtmlCSS() + df.to_html(index=False)
        # prompt = (
        #     "请将以下水库预警数据转换为直观的表格形式，表头包括：岗位、 责任人、 联系方式、 防汛职责。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   岗位    |       责任人     |   联系方式    |      防汛职责   |"
        #     "|   XXX    |        XXX      |     XXX     |       XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{xxyYingjidata}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res

def OrganizeBaoZhang_zhihuibu_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        xxyYingjidata = []
        for xxyYingjidict in xld_data['XLDheadquarters']:
            xxyYingjidata.append({'岗位':xxyYingjidict['job_desc'],'责任人':xxyYingjidict['name'],'联系方式':xxyYingjidict['concat'],'防汛职责':xxyYingjidict['duty']})
        df = pd.DataFrame(xxyYingjidata)
        res = df.to_json(orient='records', force_ascii=False)
    else:
        res = query_question(context)
    return res

def company_duty(context=None):
    if isinstance(context, dict):
        # context = json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        company_duty_data=[]
        for company_duty_dict in context['companyTable']:
            company_duty_data.append({'岗位':company_duty_dict['job_desc'],'责任人':company_duty_dict['name'],
                                      '联系方式':company_duty_dict['concat'],'防汛职责':company_duty_dict['duty']})
        df = pd.DataFrame(company_duty_data)
        res = pd2HtmlCSS() + df.to_html(index=False)
        # prompt = (
        #     "请将以下开发公司防汛指挥部成员防汛职责信息数据转换为直观的表格形式，表头包括：岗位、 责任人、 联系方式、 防汛职责。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   岗位    |       责任人     |   联系方式    |      防汛职责   |"
        #     "|   XXX    |        XXX      |     XXX     |       XXX      |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{company_duty_data}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res

def company_duty_json(context=None):
    if isinstance(context, dict):
        # context = json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))

        company_duty_data=[]
        for company_duty_dict in context['companyTable']:
            company_duty_data.append({'岗位':company_duty_dict['job_desc'],'责任人':company_duty_dict['name'],
                                      '联系方式':company_duty_dict['concat'],'防汛职责':company_duty_dict['duty']})
        df = pd.DataFrame(company_duty_data)
        res = df.to_json(orient='records', force_ascii=False)
    else:
        res = query_question(context)
    return res

def team_baozhang(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        team_baozhang_data = []
        for team_baozhang_dict in xld_data['teamTable']:
            team_baozhang_data.append({"名称":team_baozhang_dict['名称'],'责任人':team_baozhang_dict['责任人'],
            '组成':team_baozhang_dict['组成'],'人数':team_baozhang_dict['人数'],'抢险职责':team_baozhang_dict['抢险职责']})
        df = pd.DataFrame(team_baozhang_data)
        res = pd2HtmlCSS() + df.to_html(index=False)
        # prompt = (
        #     "请将以下开发公司队伍保障人员信息数据转换为直观的表格形式，表头包括：名称、 责任人、 组成、人数、抢险职责。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   名称    |       责任人     |     组成     |    人数    |    抢险职责   |"
        #     "|   XXX    |        XXX      |     XXX     |    XXX    |       XXX    |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n数据如下：{team_baozhang_data}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res

def team_baozhang_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        team_baozhang_data = []
        for team_baozhang_dict in xld_data['teamTable']:
            team_baozhang_data.append({"名称":team_baozhang_dict['名称'],'责任人':team_baozhang_dict['责任人'],
            '组成':team_baozhang_dict['组成'],'人数':team_baozhang_dict['人数'],'抢险职责':team_baozhang_dict['抢险职责']})
        df = pd.DataFrame(team_baozhang_data)
        res = df.to_json(orient='records',force_ascii=False)
    else:
        res = query_question(context)
    return res

def fangxun_table(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        fangxun_data = []
        for fangxun_dict in xld_data['goodsTable']:
            fangxun_data.append({"类别":fangxun_dict['类别'],'物资名称':fangxun_dict['名称'],
            '物资规格':fangxun_dict['品种规格'],'单位':fangxun_dict['单位'],'数量':fangxun_dict['实际库存']})
        df = pd.DataFrame(fangxun_data)
        res = pd2HtmlCSS() + df.to_html(index=False)
        # prompt = (
        #     "请将以下开发公司防汛物资信息数据转换为直观的表格形式，表头包括：类别、物资名称、 物资规格、 单位、 数量。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   类别    |      物资名称     |     物资规格     |    单位    |       数量    |"
        #     "|   XXX    |        XXX      |       XXX       |    XXX    |       XXX    |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n防汛物资信息数据数据如下：{fangxun_data}"
        # )
        # df = pd.DataFrame(fangxun_data)
        # res = tabulate(df, headers='keys', tablefmt='pretty', showindex=False)
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res
def fangxun_table_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        fangxun_data = []
        for fangxun_dict in xld_data['goodsTable']:
            fangxun_data.append({"类别":fangxun_dict['类别'],'物资名称':fangxun_dict['名称'],
            '物资规格':fangxun_dict['品种规格'],'单位':fangxun_dict['单位'],'数量':fangxun_dict['实际库存']})
        df = pd.DataFrame(fangxun_data)
        res = df.to_json(orient='records', force_ascii=False)
    else:
        res = query_question(context)
    return res

def xld_diaodu_table(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        diaodu_data = []
        for  diaodu_dict in xld_data['realDataTabel']:
            diaodu_data.append({"时间": diaodu_dict['mTime'],'库水位（m）': diaodu_dict['z'],
            '蓄水量(亿m³)': diaodu_dict['w'],'入库流量（m³/s）': diaodu_dict['qi'],'出库流量（m³/s）': diaodu_dict['qo']})
        df = pd.DataFrame(diaodu_data)
        res = pd2HtmlCSS() + df.to_html(index=False)
        #res = tabulate(df, headers='keys', tablefmt='pretty', showindex=False)
        # prompt = (
        #     "请将以下小浪底水库的调度过程数据转换为直观的表格形式，表头包括：时间、库水位（m）、蓄水量(亿m³)、 入库流量（m³/s）、 出库流量（m³/s）。"
        #     "请确保以下内容对齐并居中显示，格式如下："
        #     "|   时间    |     库水位（m）   |    蓄水量(亿m³)   | 入库流量（m³/s）|   出库流量（m³/s）   |"
        #     "|   XXX    |        XXX      |       XXX       |      XXX      |         XXX        |"
        #     "请使用空格填充以确保列宽一致，生成的表格应简洁明了，不要出现任何额外的说明性文字。"
        #     "请确保返回的只有Markdown形式的语法"
        #     f"\n防汛物资信息数据数据如下：{diaodu_data}"
        # )
        # res = query_question(prompt)
        # markdown_table = re.findall(r'(\|.*?\|(?:\n|$))', res)
        # # 使用 StringIO 读取 Markdown 表格
        # if markdown_table:  # 确保有匹配的表格
        #     # 使用 StringIO 读取 Markdown 表格
        #     data = StringIO(''.join(markdown_table))
        #
        #     # 读取为 DataFrame
        #     df = pd.read_csv(data, sep='|', skipinitialspace=True)
        #
        #     # 清洗数据，移除空列
        #     df = df.dropna(axis=1, how='all')
        #     # 使用 tabulate 输出带边框的表格
        #     final_output = tabulate(df, headers='keys', tablefmt='grid', showindex=False)
        #     res = final_output
    else:
        res = query_question(context)
    return res
def xld_diaodu_table_json(context=None):
    if isinstance(context, dict):
        xld_data = context#json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
        diaodu_data = []
        for  diaodu_dict in xld_data['realDataTabel']:
            diaodu_data.append({"时间": diaodu_dict['mTime'],'库水位（m）': diaodu_dict['z'],
            '蓄水量(亿m³)': diaodu_dict['w'],'入库流量（m³/s）': diaodu_dict['qi'],'出库流量（m³/s）': diaodu_dict['qo']})
        df = pd.DataFrame(diaodu_data)
        res =df.to_json(orient='records', force_ascii=False)
    else:
        res = query_question(context)
    return res

def levenshtein_distance(s1, s2):
    if len(s1) > len(s2):
        s1, s2 = s2, s1

    distances = range(len(s1) + 1)
    for index2, char2 in enumerate(s2):
        new_distances = [index2 + 1]
        for index1, char1 in enumerate(s1):
            if char1 == char2:
                new_distances.append(distances[index1])
            else:
                new_distances.append(1 + min((distances[index1], distances[index1 + 1], new_distances[-1])))
        distances = new_distances

    return distances[-1]

def similarity_score(s1, s2):
    lev_dist = levenshtein_distance(s1, s2)
    max_len = max(len(s1), len(s2))
    return (max_len - lev_dist) / max_len


def jaccard_similarity(s1, s2):
    # 将字符串转换为集合
    set1 = set(s1)
    set2 = set(s2)

    # 计算交集和并集
    intersection = set1.intersection(set2)
    union = set1.union(set2)

    # 计算 Jaccard 相似度
    if len(union) == 0:
        return 0.0  # 防止除以零
    return len(intersection) / len(union)


def map_input_to_label1(user_input):
    labels = [
        "雨情实况", "实时雨水情","河道水情", "水库水情", "工情险情实况", "降雨预报","河道边界条件",
        "洪水预报", "调度方案", "调度结果", "工程研判", "枢纽运用方案", "安全举措","来水预估", "调度原则和调度目标"
    ]

    user_input = user_input.lower()
    best_match = None
    best_score = 0

    for label in labels:
        label_lower = label.lower()
        score = jaccard_similarity(user_input, label_lower)
        if score > best_score:
            best_score = score
            best_match = label

    # 设置一个阈值，例如0.5，用于判断匹配的有效性
    if best_match is not None and best_score > 0.2:  # 可以调整阈值
        return best_match
    else:
        return user_input

def similarity_ratio(s1, s2):
    return difflib.SequenceMatcher(None, s1, s2).ratio()

def map_input_to_label(user_input):
    labels = [
        "雨情实况", "实时雨水情","河道水情", "水库水情", "工情险情实况", "降雨预报","河道边界条件",
        "洪水预报", "调度方案", "调度结果", "工程研判", "枢纽运用方案", "安全举措","来水预估", "调度原则和调度目标"
    ]

    user_input = user_input.lower()
    best_match = None
    best_score = 0

    for label in labels:
        label_lower = label.lower()
        score = similarity_ratio(user_input, label_lower)
        if score > best_score:
            best_score = score
            best_match = label

    # 设置一个阈值，例如0.5，用于判断匹配的有效性
    if best_match is not None and best_score > 0.2:  # 可以根据实际情况调整阈值
        return best_match
    else:
        return user_input

def recommend_plan(text, plans: List[PlanTemplate] = [], user: User=None):
    """
        基于用户行为, 推荐合适的模板, 相似打分
    """
    results = []
    # if user:
    #     # TODO 实现基于历史推荐算法
    #     for plan in plans:
    #         results.append(model_to_dict(plan, exclude=['nodes']))
    #     return results
    # else:
    if text:
        # TODO 关键词匹配算法  优化
        for plan in plans:
            if plan.name.find(text) != -1:
                tmp = model_to_dict(plan, exclude=[])
                tmp['keywords'] = text
                tmp['score'] = 1
                results.append(model_to_dict(plan, exclude=['nodes']))
    else:
        # 默认推荐
        for plan in plans:
            tmp = model_to_dict(plan, exclude=[])
            tmp['keywords'] = text
            tmp['score'] = 0.0
            results.append(model_to_dict(plan, exclude=['nodes']))
    return results

def map_input_to_template(user_input, templates):
    user_input = user_input.lower()
    best_match = None
    best_score = 0

    for label in templates:
        label_lower = label.lower()
        score = similarity_ratio(user_input, label_lower)
        if score > best_score:
            best_score = score
            best_match = label

    # 设置一个阈值，例如0.5，用于判断匹配的有效性
    if best_match is not None and best_score > 0.2:  # 可以根据实际情况调整阈值
        return best_match
    else:
        return user_input

def generate_description_for_label(label):
    frag = search_fragpacks(label)
    if frag:
        prompt = f"请围绕'{label}', 依据{frag},撰写一段相关的文字描述并进行优化，结果要凝练简洁，不要出现冗余描述，写的专业化些。"
    else:
        prompt = f"请围绕'{label}', 撰写一段相关的文字描述并进行优化，结果要凝练简洁，不要出现冗余描述，写的专业化些。"
    description = query_question(prompt)
    return description


def xld_safety_measure(context):
    xld_yingji=YingjiResponse(context)
    zuzhibaozhang = OrganizeBaoZhang_leader(context)
    zhihuibaozhang = OrganizeBaoZhang_zhihuibu(context)
    gongsi_duty = company_duty(context)
    duiwu_baozhang = team_baozhang(context)
    fangxun_information = fangxun_table(context)
    diaodu_table = xld_diaodu_table(context)
    safety_result = (f"四   安全举措\n"
                           f"4.1  应急响应措施\n"
                           f"{xld_yingji}\n"
                           f"4.2  应急保障\n"
                           f"组织保障\n"
                           f"小浪底管理中心防汛领导小组人员信息表\n"
                           f"{zuzhibaozhang}\n"
                           f"小浪底水利枢纽、西霞院工程防汛应急抢险指挥部\n"
                           f"{zhihuibaozhang}\n"
                           f" 开发公司防汛指挥部成员防汛职责信息表\n"
                           f"{gongsi_duty}\n"
                           f"队伍保障人员信息表\n"
                           f"{duiwu_baozhang}\n"
                           f"防汛物资信息表\n"
                           f"{fangxun_information}\n"
                           f"附表1\n 预报小浪底水库的调度过程\n"
                           f"{diaodu_table}"
                     )
    return safety_result


def engineer_safety_xldddp(context=None):
    if context is None:
        context = json.load(open("data/qiuxun.json", mode='r', encoding="utf8"))
    ybddgucTab = []
    for ent in context['realDataTabel']:
        ybddgucTab.append({"mTime": ent['mTime'], "z": ent['z'], "w": ent['w'], "qi": ent['qi'], "qo": ent['qo']})
    context['ddTab'] = ybddgucTab
    ddTab_analysis = []
    for ent in context['realDataTabel']:
        ddTab_analysis.append({"时间": ent['mTime'], "入库流量（m³/s）": ent['z'], "出库流量（m³/s）": ent['w'], "蓄水量(亿m³)": ent['qi'], "库水位（m）": ent['qo']})
    context['ddTab_analysis'] = ddTab_analysis

    # prompt = (f"请将以下列表绘制成表格，表头为：时间、库水位（m）、蓄水量（亿m³）、入库流量（m³/s）、出库流量（m³/s）。表格名称为“预报小浪底水库的调度过程”。"
    #           f"\n列表内容如下：{ybddgucTab}")
    prompt = (f"请将以下列表绘制成表格，表头为：时间、库水位（m）、蓄水量（亿m³）、入库流量（m³/s）、出库流量（m³/s）。"
              f"表格名称为“预报小浪底水库的调度过程”。"
              f"请确保所有内容居中对齐, 并确保是以表格形式输出。"
              f"\n列表内容如下：{ddTab_analysis}")
    # prompt = (f"请从以下列表中提取关键信息，包括“水库水位变化”、“蓄水量信息”和“流量信息”，并进行总结"
    #           f"\n列表内容如下：{ybddgucTab}")
    llm = Ollama(model="qwen2")
    res = llm(prompt)
    return res

def generate_docx(context=None):
    if context is None:
        context = {
        "一级标题": {
            "一、雨水情势": [],
            "二、安全研判": [],
            "三、枢纽运用": ["3.1 调度过程", "3.2 运用推荐"],
            "四、安全举措": ["4.1 响应措施", "4.2 应急保障"]
        }
    }
        
    template = DocxTemplate(r"data\template.docx")  # 假设你有一个模板文件data/template.docx
    template.render(context)
    # 保存生成的文档
    template.save("output.docx")
    # doc = Document("output.docx")
    #
    # for para in doc.paragraphs:
    #     # 在每个段落后添加换行
    #     para.add_run().add_break()  # 添加换行
    #
    # doc.save("output.docx")
    # res = yushuiqing()
def shj_yushui_context(context=None):
    default_context = "雨情：6月22日，兰州以上部分，兰托区间局部地区降小雨，个别站中雨；黄河流域其他地区降分散性小雨。最大点雨量出现在黄河上游凉坪站和石骨岔站，均为15毫米。"
    now = datetime.now()
        # 获取当前的年、月和日
    year = now.year
    month = now.month
    day = now.day
    if isinstance(context, dict):
        huanghe_weather = context.get("huanghe_weather")  # 从 context 中获取值
        lanzhou_weather = context.get("lanzhou_weather")
        yuliang = context.get("yuliang")
        information = (
            f"兰州区域天气情况：{lanzhou_weather}，"
            f"黄河流域天气情况：{huanghe_weather}，"
            f"降雨量情况：{yuliang}"
        )
        prompt = (
            f"参考描述：{default_context}\n"
            f"请模仿上述描述，根据以下已知信息生成雨水实况，并进行优化，不要生成无关信息提示，"
            f"也请不要使用'优化后的描述：'这样的字样。"
            f"确保不包含任何说明性文字。"
            f"\n已知信息：{information},当前日期: {year}年{month}月{day}日"
        )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def shj_shuikuxushui_generate_dfjson(context):
    if isinstance(context, dict):
        default_list = context['sksq']
        df = pd.DataFrame(default_list)
        res = df.to_json(orient='records')
    else:
        res = query_question(context)
    return res

def shj_shuikuxushui_generate(context):
    if isinstance(context, dict):
        default_list = context['sksq']
        df = pd.DataFrame(default_list)
        res = pd2HtmlCSS() + divHtml(df.to_html(index=False))
    else:
        res = query_question(context)
    return res

def shj_hedaoshuiqing_generate_dfjson(context):
    if isinstance(context, dict):
        default_list = context['hdsq']
        df = pd.DataFrame(default_list)
        res = df.to_json(orient='records')
    else:
        res = query_question(context)
    return res

def shj_hedaoshuiqing_generate(context):
    if isinstance(context, dict):
        default_list = context['hdsq']
        df = pd.DataFrame(default_list)
        res = pd2HtmlCSS() + divHtml(df.to_html(index=False))
    else:
        res = query_question(context)
    return res

def shj_laishuiyugu_context(context=None):
    default_context = "考虑未来降雨，预估未来七天，唐乃亥站日均流量为1240～1300m3/s，总水量约7.63亿m3；龙刘区间、刘兰区间水量分别约为1.25亿m3和1.14亿m3。头道拐站日均流量为370～510m3/s，潼关站日均流量为470～630m3/s。"
    if isinstance(context, dict):
        tng_rjll = context.get("tng_rjll")  # 从 context 中获取值
        tng_zhl = context.get("tng_zhl")
        llqujian = context.get("llqujian")
        llqj = context.get("llqj")
        tgd_rjll = context.get("tgd_rjll")
        tg_rjll = context.get("tg_rjll")
        information = (
            f"唐乃亥日均流量：{tng_rjll}，"
            f"唐乃亥总水量：{tng_zhl}，"
            f"龙刘区间流量：{llqujian}，"
            f"刘兰区间流量：{llqj}，"
            f"头道拐日均流量：{tgd_rjll}，"
            f"潼关日均流量：{tg_rjll}"
        )
        prompt = (
            f"参考描述：{default_context}\n"
            f"请模仿上述描述，根据以下已知信息生成雨水实况，并进行优化，不要生成无关信息提示，"
            f"也请不要使用'优化后的描述：'这样的字样。"
            f"确保不包含任何说明性文字。"
            f"\n已知信息：{information}"
        )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res

def shj_laishuiyugu_generate_dfjson(context):
    if isinstance(context, dict):
        default_list = context['lsyg']
        df = pd.DataFrame(default_list)
        res = df.to_json(orient='records')
    else:
        res = query_question(context)
    return res

def shj_laishuiyugu_generate(context):
    if isinstance(context, dict):
        default_list = context['lsyg']
        df = pd.DataFrame(default_list)
        res = pd2HtmlCSS() + divHtml(df.to_html(index=False))
    else:
        res = query_question(context)
    return res

def shj_hedaobijian_context(context=None):
    default_context = "据分析预估，黄河上游龙羊峡水库以下河道安全过洪流量，青海、甘肃、宁夏、内蒙古河段分别约为3660m3/s、3730m3/s、5620m3/s、5510m3/s。黄河中游小北干流河道最小平滩流量约为3000m3/s；黄河下游河道最小平滩流量为4600m3/s。"
    if isinstance(context, dict):
        qh_aqghll = context.get("qh_aqghll")  # 从 context 中获取值
        gs_aqghll = context.get("gs_aqghll")
        nx_aqghll = context.get("nx_aqghll")
        nm_aqghll = context.get("nm_aqghll")
        xbgl_aqghll = context.get("xbgl_aqghll")
        hhxx_aqghll = context.get("hhxx_aqghll")
        information = (
            f"青海河道安全过洪流量：{qh_aqghll}，"
            f"甘肃河道安全过洪流量：{gs_aqghll}，"
            f"宁夏河道安全过洪流量：{nx_aqghll}，"
            f"内蒙古河道安全过洪流量：{nm_aqghll}，"
            f"小北干流河道最小平滩流量：{xbgl_aqghll}，"
            f"黄河下游河道最小平滩流量：{hhxx_aqghll}"
        )
        prompt = (
            f"参考描述：{default_context}\n"
            f"请模仿上述描述，根据以下已知信息生成雨水实况，并进行优化，不要生成无关信息提示，"
            f"也请不要使用'优化后的描述：'这样的字样。"
            f"确保不包含任何说明性文字。"
            f"\n已知信息：{information}"
        )
        res = query_question(prompt)
    else:
        res = query_question(context)
    return res
# res = huanghe_yuqing_generate(context="黄河中下游洪水调度方案")
# print(res)

#res = huanghe_hedaoshuiqing_generate(context="帮我生成黄河中下游洪水调度方案")
#res = huanghe_shuikushuiqing_generate(context="帮我生成黄河中下游洪水调度方案")
#res = huanghe_gongqing_generate(context="帮我生成黄河中下游洪水调度方案")
#res = realtime_yushuigongqing(context="帮我生成黄河中下游实时雨水工险情")


# res = huanghe_jiangyu13_forecast(context="帮我生成黄河中下游实时雨水工险情")
# res = huanghe_jiangyu47_forecast(context="帮我生成黄河中下游实时雨水工险情")
#res = huanghe_fenqu_jiangyu_forecast(context="帮我生成黄河中下游实时雨水工险情")
#res = huanghe_flood_forecast(context="帮我生成黄河中下游实时雨水工险情")
# res = huanghe_yubao(context="帮我生成黄河中下游实时雨水工险情")


# res = huanghe_diaodu_plan(context="帮我生成黄河中下游实时雨水工险情")
# res = huanghe_diaodu_table(context="帮我生成黄河中下游实时雨水工险情")
# res = huanghe_diaodu(context="帮我生成黄河中下游实时雨水工险情")


# res = huanghe_shuiku_diaodu_result(context="帮我生成黄河中下游实时雨水工险情")
#res = huanghe_tanqu_yanmo(context="帮我生成黄河中下游实时雨水工险情")
#res = huanghe_keneng_danger(context="帮我生成黄河中下游实时雨水工险情")
#res = huanghe_xiangying_level(context="帮我生成黄河中下游实时雨水工险情")
# res = huanghe_diaodu_result(context="帮我生成黄河中下游实时雨水工险情")
# print(res,'\n\n')

# res =xld_yushui_context("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = engineer_safety_shuikuyj("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = engineer_safety_shuiwenyj("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = engineer_safety_gongchengjcyj("帮我生成小浪底防汛预案")
# res = xld_engineer_safety("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)

# res = shuniuFangAn(context="帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = xldshuiddProcessAnalysys()
# print("engineer-safety:",res)
# res = xldJZStatus(context="帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = xldholeStatus(context="帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res =JZHoleRecommend(context="帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = xld_shuniu_apply(context="帮我生成小浪底防汛预案")
# print("engineer-safety:",res)


# res = YingjiResponse("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = OrganizeBaoZhang_leader("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = OrganizeBaoZhang_zhihuibu("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = OrganizeBaoZhang()
# print("engineer-safety:",res)
# res = company_duty("帮我生成小浪底防汛预案")
# print("engineer-safety:",res)
# res = team_baozhang()
# print("engineer-safety:",res)
# res = fangxun_table()
# print("engineer-safety:",res)
# res = engineer_safety_xldddp()
# print("engineer-safety:",res)
# generate_docx()

#qiuxun2Word()


# @app.get("/")
# async def document():
#     return RedirectResponse(url="/docs")

# @app.get("/realtime_yushuigongqing")
# async def get_yushuigongqing(context: str):
#     try:
#         result = realtime_yushuigongqing(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
# #
# @app.get("/huanghe_yubao")
# async def get_huanghe_yubao(context: str):
#     try:
#         result = huanghe_yubao(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


# @app.get("/huanghe_diaou")
# async def gethuanghe_diaodu(context: str):
#     try:
#         result = huanghe_diaodu(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.get("/huanghe_diaou_reulst")
# async def get_huanghe_diaodu_result(context: str):
#     try:
#         result = huanghe_diaodu_result(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.get("/xld_yushui")
# async def get_xld_yushui_result(context: str):
#     try:
#         result =xld_yushuqing(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


# @app.get("/xld_engineer_safety")
# async def get_xld_engineer_safety(context: str):
#     try:
#         result = xld_engineer_safety(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.get("/xld_shuniu_apply")
# async def get_xld_engineer_safety(context: str):
#     try:
#         result = xld_shuniu_apply(context)
#         return {"code": 200, "data": result}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


# if __name__ == "__main__":
#     import uvicorn

#     app.add_middleware(
#         CORSMiddleware,
#         allow_origins=["*"],  # 允许所有来源
#         allow_credentials=True,
#         allow_methods=["*"],  # 允许所有HTTP方法
#         allow_headers=["*"],  # 允许所有请求头
#     )
#     uvicorn.run(app, host="0.0.0.0", port=7777)

