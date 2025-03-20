import sys
from typing import List
from django.shortcuts import render
import json
import pprint
import time
from django.shortcuts import render
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
import pandas as pd
from rest_framework.viewsets import ViewSet, ModelViewSet
from rest_framework.views import APIView
from rest_framework.decorators import action
from django.http.multipartparser import MultiPartParser
from rest_framework.response import Response
from django.http import FileResponse, HttpResponse, JsonResponse, StreamingHttpResponse
# from django.contrib.auth.models import User, Group
from rest_framework import status
from rest_framework import mixins
from rest_framework import generics
from django.forms.models import model_to_dict
from django.db.models import Q  
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.views.decorators.csrf import csrf_exempt
from datetime import date, datetime
from rest_framework.authentication import BasicAuthentication
# 准备数据
from pyecharts.charts import Line
from pyecharts import options as opts
import imgkit
import os
import collections
from docx.enum.section import WD_SECTION
from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.files.base import ContentFile
from rest_framework.decorators import api_view
from docx import Document  

from rest_framework.parsers import (
    FormParser,
    MultiPartParser
)

from kgapp.models import KgBusiness, KgProductTask, KgTask
from userapp.models import User
from yaapp import yautils
from yaapp import getYuAnName, getYuAnParamPath
from yaapp.api_yuan import map_input_to_template, recommend_plan
from yaapp.models import PlanByUser, PlanByUserDocument, PlanTemplate, PtBusiness, TemplateNode
from yaapp.plan import PlanFactory
from yaapp.wordutils import set_landscape, writeParagraphs2Word, writeTitle2Word
from yunheKGPro import CsrfExemptSessionAuthentication

from yaapp.serializers import YuAnAppResponseSerializer


class PTBusinessList(mixins.ListModelMixin,
                  mixins.CreateModelMixin,
                  generics.GenericAPIView):
    
    serializer_class = YuAnAppResponseSerializer
    @swagger_auto_schema(
            operation_description='GET /ptbusinesslist',
            operation_summary="预案的业务类型",
            # 接口参数 GET请求参数
            manual_parameters=[
            ],
            responses={
                200: YuAnAppResponseSerializer(many=False),
                400: "请求失败",
            },
            tags = ['ya_api'])
    def get(self, request, *args, **kwargs):
        data = {"code": 200}
        bus_list = PtBusiness.objects.all()
        data['data'] = [model_to_dict(b) for b in bus_list]
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data,  status=status.HTTP_200_OK)


class BlockList(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="[板块列表] 获取预案模板节点列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('keyword', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('page', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
            openapi.Parameter('pageSize', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        keyword = request.GET.get("keyword", None)
        page = request.GET.get("page", 1)
        pageSize = request.GET.get("pageSize", 5)
        BlockList.queryset = TemplateNode.objects

        if keyword:
            BlockList.queryset = BlockList.queryset.filter(label__icontains=keyword).order_by('-created_at').all()
        else:
            BlockList.queryset = BlockList.queryset.order_by('-created_at').all()

        paginator = Paginator(BlockList.queryset, pageSize)
        try:
            ptList = paginator.page(page)
        except PageNotAnInteger:
            ptList = paginator.page(1)
        except EmptyPage:
            ptList = paginator.page(paginator.num_pages)

        results = []
        for tmp in ptList:
            tmpdict = model_to_dict(tmp, exclude=["wordParagraphs", 'result'])
            tmpdict['created_at'] = tmp.created_at.strftime("%Y-%m-%d %H:%M:%S")
            results.append(tmpdict)

        data = {"code": 200, "msg": "success", "data": results, "success": True, "total": paginator.count, "page": paginator.num_pages, "pageSize": paginator.per_page}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)

class BlockOptionList(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="[板块配置] 选择可用的板块列表",
        # 接口参数 GET请求参数
        manual_parameters=[
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        BlockOptionList.queryset = TemplateNode.objects
        BlockOptionList.queryset = BlockOptionList.queryset.values('label').distinct()
        results = [ent.get('label') for ent in BlockOptionList.queryset]
        data = {"code": 200, "msg": "success", "data": results}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)
    
# Create your views here.
class TemplateList(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="[模板列表] 获取预案模板列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('keyword', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('page', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
            openapi.Parameter('pageSize', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        keyword = request.GET.get("keyword", None)
        page = request.GET.get("page", 1)
        pageSize = request.GET.get("pageSize", 5)
        TemplateList.queryset = PlanTemplate.objects

        if keyword:
            TemplateList.queryset = TemplateList.queryset.filter(name__icontains=keyword).order_by('-created_at').all()
        else:
            TemplateList.queryset = TemplateList.queryset.order_by('-created_at').all()

        paginator = Paginator(TemplateList.queryset, pageSize)
        try:
            ptList = paginator.page(page)
        except PageNotAnInteger:
            ptList = paginator.page(1)
        except EmptyPage:
            ptList = paginator.page(paginator.num_pages)

        results = []
        for tmp in ptList:
            tmpdict = model_to_dict(tmp, exclude=["nodes"])
            tmpdict['created_at'] = tmp.created_at.strftime("%Y-%m-%d %H:%M:%S")
            results.append(tmpdict)

        data = {"code": 200, "msg": "success", "data": results, "success": True, "total": paginator.count, "page": paginator.num_pages, "pageSize": paginator.per_page}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)
    

class TemplateDetail(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="获取预案模板列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('id', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        ptId = request.GET.get("id", 1)
        try:
            pt = PlanTemplate.objects.get(id=ptId)
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
            bars = YuAnAppResponseSerializer(data=data, many=False)
            bars.is_valid()
            return Response(bars.data, status=status.HTTP_200_OK)

        tmpResult = model_to_dict(pt, exclude=["nodes"])
        tmpResult['node_list'] = pt.nodelist
        tmpResult['created_at'] = pt.created_at.strftime("%Y-%m-%d %H:%M:%S")
        data = {"code": 200, "msg": "success", "data": tmpResult, "success": True}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)
    
class AddOrUpdateTemplate(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='POST 节点或者更新预案',
        operation_description='节点或者更新预案',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_INTEGER, description="ID"),
                'name': openapi.Schema(type=openapi.TYPE_STRING, description="预案名称"),
                'description': openapi.Schema(type=openapi.TYPE_STRING, description="预案相关描述"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print("TemplateAddOrUpdateNode:", request.data)
        ptid = request.data.get("id", None)
        name = request.data.get("name", "")
        description = request.data.get("description", "")

        if ptid is None:
            # create
            # order = parent.node_cnt if parent is not None else 0
            tmpPT = PlanTemplate.objects.create(name=name, description=description)
            tmpResult = model_to_dict(tmpPT, exclude=["nodes"])
            tmpResult['node_list'] = tmpPT.nodelist
            data = {"code": 200, "data": tmpResult, "msg": "新预案模板创建成功！"}
        else:
            try:
                tmpPT = PlanTemplate.objects.get(id=ptid)
                if name:
                    tmpPT.name = name
                if description:
                    tmpPT.description = description
                tmpPT.save()
                tmpResult = model_to_dict(tmpPT, exclude=["nodes"])
                tmpResult['node_list'] = tmpPT.nodelist
                data = {"code": 200, "data": tmpResult, "msg": "预案更新成功！"}
            except:
                data = {"code": 201, "data": {}, "msg": "参数错误"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    
class DeleteTemplate(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer
    @swagger_auto_schema(
        operation_summary='删除预案操作',
        operation_description='POST ///',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_INTEGER, description="ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        ptId = request.data.get("id", None)
        try:
            pt = PlanTemplate.objects.get(id=ptId)
            pt.delete()
            data = {"code": 200, "data": {}, "msg": "删除成功！", "success": True}
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}

        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    


class AddOrUpdateTemplateByNode(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='POST',
        operation_description='节点或者更新预案',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'ptid': openapi.Schema(type=openapi.TYPE_INTEGER, description="预案ID"),
                'ntid': openapi.Schema(type=openapi.TYPE_INTEGER, description="节点ID"),
                'label': openapi.Schema(type=openapi.TYPE_STRING, description="label"),
                'order': openapi.Schema(type=openapi.TYPE_INTEGER, description="order"),
                'description': openapi.Schema(type=openapi.TYPE_STRING, description="description"),
                'template': openapi.Schema(type=openapi.TYPE_STRING, description="template"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print("AddOrUpdateTemplateByNode:", request.data)
        
        ptid = request.data.get("ptid", None)
        ntid = request.data.get("ntid", None)
        label = request.data.get("label", "")
        order = request.data.get("order", 0)
        description = request.data.get("description", "")
        template = request.data.get("template", "")

        if ptid is None:
            data = {"code": 201, "data": {}, "msg": "参数错误"}        
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        else:
            if ntid is None:
                # 创建节点
                tmpnode = TemplateNode.objects.create(label=label, description=description, template=template, order=order)
                tmpPT = PlanTemplate.objects.get(id=ptid)
                tmpPT.nodes.add(tmpnode)
                tmpPT.save()
                tmpResult = model_to_dict(tmpPT, exclude=["wordParagraphs", "result"])
                data = {"code": 200, "data": tmpResult, "msg": "节点创建成功！"}
            else:
                # 更新节点
                try:
                    nt = TemplateNode.objects.get(id=ntid)
                    if label:
                        nt.label = label
                    if description:
                        nt.description = description
                    if template:
                        nt.template = template
                    if order:
                        nt.order = order
                    nt.save()
                    tmpResult = model_to_dict(nt, exclude=["wordParagraphs", "result"])
                    data = {"code": 200, "data": tmpResult, "msg": "节点更新成功！"}
                except:
                    data = {"code": 201, "data": {}, "msg": "参数错误"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    

class CreateSysTemplate(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='POST 【后台管理】新建模板',
        operation_description='【后台管理】新建模板',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=[],
            properties={
                'name': openapi.Schema(type=openapi.TYPE_STRING, description="name"),
                'description': openapi.Schema(type=openapi.TYPE_STRING, description="description"),
                'created_at': openapi.Schema(type=openapi.TYPE_STRING, description="created_at"),
                'businessid': openapi.Schema(type=openapi.TYPE_NUMBER, description="businessid 业务ID"),
                # 'blockOpts': openapi.Schema(type=openapi.TYPE_ARRAY, items=openapi.Schema(type=openapi.TYPE_STRING, properties={}), description="blackids 涉及板块ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print("AddOrUpdateTemplateByNode:", request.data)
        
        name = request.data.get("name", None)
        description = request.data.get("description", "")
        created_at = request.data.get("created_at", "")
        businessid = request.data.get("businessid", None)
        blockOpts = request.data.get("blockOpts", [])

        if not name or not businessid or not created_at or not blockOpts:
            data = {"code": 201, "data": {}, "msg": "参数错误"}        
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        else:
            tmpPT = PlanTemplate.objects.create(name=name, description=description, created_at=created_at, ctype=businessid, published=False)
            for i, opt in enumerate(blockOpts):
                upload_file_flag = False
                if opt == "调度方案":
                    upload_file_flag = True
                bd = TemplateNode.objects.create(label=opt, description="", template="", order= i+1, upload_file_flag=upload_file_flag)
                tmpPT.nodes.add(bd)
            tmpPT.save()
            retJson = model_to_dict(tmpPT, exclude=["nodes"])
            retJson['nodeOutlineList'] = tmpPT.nodeOutlineList
            data = {"code": 200, "data": retJson, "msg": "节点创建成功！"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)


class UpdateTemplateNode(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='POST 【板块配置信息】',
        operation_description='【板块配置信息】',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'nodeid': openapi.Schema(type=openapi.TYPE_INTEGER, description="节点ID"),
                'label': openapi.Schema(type=openapi.TYPE_STRING, description="label"),
                'order': openapi.Schema(type=openapi.TYPE_INTEGER, description="order"),
                'description': openapi.Schema(type=openapi.TYPE_STRING, description="description"),
                'template': openapi.Schema(type=openapi.TYPE_STRING, description="template"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print("AddOrUpdateTemplateByNode:", request.data)
        
        nodeid = request.data.get("nodeid", None)
        label = request.data.get("label", "")
        order = request.data.get("order", 0)
        description = request.data.get("description", "")
        template = request.data.get("template", "")

        if nodeid is None:
            data = {"code": 201, "data": {}, "msg": "参数错误"}        
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        else:
            try:
                nt = TemplateNode.objects.get(id=nodeid)
                if label:
                    nt.label = label
                if description:
                    nt.description = description
                if template:
                    nt.template = template
                if order:
                    nt.order = order
                nt.save()
                data = {"code": 200, "data": model_to_dict(nt, exclude=["wordParagraphs"]), "msg": "节点更新成功！"}
            except:
                data = {"code": 201, "data": {}, "msg": "参数错误"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)

class DeleteNodeByTemplate(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer
    @swagger_auto_schema(
        operation_summary='删除预案操作',
        operation_description='POST ///',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'ptid': openapi.Schema(type=openapi.TYPE_INTEGER, description="预案ID"),
                'ntid': openapi.Schema(type=openapi.TYPE_INTEGER, description="节点ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        ptId = request.data.get("ptid", None)
        ntId = request.data.get("ntid", None)

        if ptId is None or ntId is None:
            data = {"code": 201, "data": {}, "msg": "参数错误"}
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)

        try:
            pt = PlanTemplate.objects.get(id=ptId)
            nt = TemplateNode.objects.get(id=ntId)
            pt.nodes.remove(nt)
            pt.save()
            data = {"code": 200, "data": {}, "msg": "删除成功！", "success": True}
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}

        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    

class LLMSingleNodePlan(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='单章节节点预案生成',
        operation_description='POST 单节点预案生成',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'ptid': openapi.Schema(type=openapi.TYPE_INTEGER, description="ptID 预案ID"),
                'nodeid': openapi.Schema(type=openapi.TYPE_INTEGER, description="nodeID 章节ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print("params:", request.data)
        nodeId = request.data.get("nodeid", None)
        planId = request.data.get("ptid", None)

        userYuAnPlan = PlanByUser.objects.get(id=planId)
        node = TemplateNode.objects.get(id=nodeId)
        print("planTemp:", userYuAnPlan, "\n", node)
        # 通用生成方法
        # node.result = qwty(node.description)
        tmp_param_path = getYuAnParamPath(userYuAnPlan.ctype, userYuAnPlan.yadate)
        if not os.path.exists(tmp_param_path):
            data = {"code": 201, "data": {}, "msg": "参数文件不存在, 请先搜集参数"}
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        
        ctx = {
            "type": userYuAnPlan.ctype,
            "yadate": userYuAnPlan.yadate,
            "plan": model_to_dict(userYuAnPlan, exclude=["html_data", "html_data", "created_at", "updated_at", "nodes"]),
            "param_path": tmp_param_path
        }
        print("PlanFactory ctx:", ctx)
        # 制作预案工厂类
        pf = PlanFactory(context=ctx, node=node)
        # 生成对应描述
        pf.make_context()
        data = {"code": 200, "data": model_to_dict(node, exclude=['parent', 'wordParagraphs']), "msg": "生成成功！"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    

class RagSingleNodePlan(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='单章节节点预案生成',
        operation_description='POST 单节点预案生成',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'ptid': openapi.Schema(type=openapi.TYPE_INTEGER, description="ptID 预案ID"),
                'nodeid': openapi.Schema(type=openapi.TYPE_INTEGER, description="nodeID 章节ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print("params:", request.data)
        nodeId = request.data.get("nodeid", None)
        planId = request.data.get("ptid", None)

        userYuAnPlan = PlanByUser.objects.get(id=planId)
        node = TemplateNode.objects.get(id=nodeId)
        print("planTemp:", userYuAnPlan, "\n", node)
        # 通用生成方法
        # node.result = qwty(node.description)
        tmp_param_path = getYuAnParamPath(userYuAnPlan.ctype, userYuAnPlan.yadate)
        if not os.path.exists(tmp_param_path):
            data = {"code": 201, "data": {}, "msg": "参数文件不存在, 请先搜集参数"}
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        
        ctx = {
            "type": userYuAnPlan.ctype,
            "yadate": userYuAnPlan.yadate,
            "plan": model_to_dict(userYuAnPlan, exclude=["html_data", "html_data", "created_at", "updated_at", "nodes"]),
            "param_path": tmp_param_path
        }
        print("PlanFactory ctx:", ctx)
        # 制作预案工厂类
        pf = PlanFactory(context=ctx, node=node)
        # 生成对应描述
        pf.make_context()
        data = {"code": 200, "data": model_to_dict(node, exclude=['parent', 'wordParagraphs']), "msg": "生成成功！"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    

class LLMNodePlan(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='预案生成',
        operation_description='POST 全文预案生成',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'ptid': openapi.Schema(type=openapi.TYPE_INTEGER, description="ptID 预案ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        #TODO 预案生成
        planId = request.data.get("ptid", None)
        print("params:", request.data)
        userYuAnPlan = PlanByUser.objects.get(id=planId)
        print("userYuAnPlan:", userYuAnPlan)
        # 通用生成方法
        # node.result = qwty(node.description)
        tmp_param_path = getYuAnParamPath(userYuAnPlan.ctype, userYuAnPlan.yadate)
        print("tmp_param_path:", tmp_param_path)
        if not os.path.exists(tmp_param_path):
            data = {"code": 201, "data": {}, "msg": "参数文件不存在, 请先搜集参数"}
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        
        ctx = {
            "type": userYuAnPlan.ctype,
            "yadate": userYuAnPlan.yadate,
            "plan": model_to_dict(userYuAnPlan, exclude=["html_data", "html_data", "created_at", "updated_at", "nodes"]),
            "param_path": tmp_param_path
        }
        for node in userYuAnPlan.nodes.all():
            # 制作预案工厂类
            pf = PlanFactory(context=ctx, node=node)
            # 生成对应描述
            pf.make_context()

        resultJson =  model_to_dict(userYuAnPlan, exclude=['parent', "nodes"])
        resultJson['nodeList'] = userYuAnPlan.nodeDetailList
        data = {"code": 200, "data": resultJson, "msg": "生成成功！"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)

class MakePlanWord(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='POST 生成预案word',
        operation_description='POST 预案写入word文件',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_INTEGER, description="ID"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print(request.data)
        pid = request.data.get("id", None)
        if pid is None:
            data = {"code": 201, "msg": "参数错误"}
            bars = YuAnAppResponseSerializer(data=data, many=False)
            bars.is_valid()
            return Response(bars.data, status=status.HTTP_200_OK)

        tmpTemplate = PlanByUser.objects.get(id=pid)
        # 创建一个Word文档对象
        time = datetime.now().strftime("%Y%m%d%H%M%S")
        doc = Document()
        set_landscape(doc)
        writeTitle2Word(tmpTemplate.name, doc)
        paragraph_format = doc.styles['Heading 1'].paragraph_format
        # 设置段落编号
        # paragraph_format.numeration = True
        for i, node in enumerate(tmpTemplate.nodelist):
            print(i, "node:", node)
            writeParagraphs2Word(i, node, doc)
        # 保存文档 
        doc.save(f'media/plans/{tmpTemplate.name}.docx')
        tmpdoc = PlanByUserDocument.objects.create(document=f'plans/{tmpTemplate.name}.docx', name=tmpTemplate.name)
        tmpdict = model_to_dict(tmpdoc, exclude=["plan", "document"])
        tmpdict['created_at'] = tmpdoc.created_at.strftime("%Y-%m-%d %H:%M:%S")
        tmpdict['file_path'] = tmpdoc.document.url
        data = {"code": 200, "data": tmpdict, "msg": "生成Word成功！"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)

class UserTemplateDocumentList(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="获取用户预案模板列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('keyword', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('page', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
            openapi.Parameter('pageSize', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        keyword = request.GET.get("keyword", None)
        page = request.GET.get("page", 1)
        pageSize = request.GET.get("pageSize", 5)
        UserTemplateDocumentList.queryset = PlanByUserDocument.objects

        if keyword:
            UserTemplateDocumentList.queryset = UserTemplateDocumentList.queryset.filter(name__icontains=keyword).order_by('-created_at').all()
        else:
            UserTemplateDocumentList.queryset = UserTemplateDocumentList.queryset.order_by('-created_at').all()

        paginator = Paginator(UserTemplateDocumentList.queryset, pageSize)
        try:
            ptList = paginator.page(page)
        except PageNotAnInteger:
            ptList = paginator.page(1)
        except EmptyPage:
            ptList = paginator.page(paginator.num_pages)

        results = []
        for tmp in ptList:
            tmpdict = model_to_dict(tmp, exclude=["plan", "document"])
            tmpdict['created_at'] = tmp.created_at.strftime("%Y-%m-%d %H:%M:%S")
            tmpdict['file_path'] = tmp.document.url
            results.append(tmpdict)
        print(results)
        data = {"code": 200, "msg": "success", "data": results, "success": True, "total": paginator.count, "page": paginator.num_pages, "pageSize": paginator.per_page}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)


import urllib.parse

def url_to_chinese(url):
    # 解析URL，提取路径部分
    path = urllib.parse.urlparse(url).path
    
    # 找到路径中的URL编码部分
    encoded_parts = path.split('/')
    chinese_part = None

    pathTokens = []
    for part in encoded_parts:
        if '%' in part:
            decoded_part = urllib.parse.unquote(part)
            pathTokens.append(decoded_part) 
        else:
            pathTokens.append(part)    
    return "/".join(pathTokens)


class DeleteUserPlanDocument(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer
    @swagger_auto_schema(
        operation_summary='删除用户预案文档操作',
        operation_description='POST ///',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_INTEGER, description="预案文档列表"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        ptDocId = request.data.get("id", None)

        if ptDocId is None:
            data = {"code": 201, "data": {}, "msg": "参数错误"}
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)

        try:
            pt = PlanByUserDocument.objects.get(id=ptDocId)
            tmpdoc = pt.document
            file_path = tmpdoc.url[1:]
            # 示例URL
            file_path = url_to_chinese(file_path)
            print(file_path)
            try:
                os.remove(file_path)
                print(f"文件 {file_path} 已被删除。")
            except FileNotFoundError:
                print(f"文件 {file_path} 不存在，无法删除。")
            except PermissionError:
                print(f"没有权限删除文件 {file_path}。")
            except Exception as e:
                print(f"删除文件时发生错误: {e}")
            
            pt.delete()
            data = {"code": 200, "data": {}, "msg": "删除成功！", "success": True}
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}

        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    


class UpdateUserPlanDocument(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer
    @swagger_auto_schema(
        operation_summary='更新用户预案文档操作',
        operation_description='POST ///',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_INTEGER, description="预案文档列表"),
                'name': openapi.Schema(type=openapi.TYPE_STRING, description="预案文档名称"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        ptDocId = request.data.get("id", None)
        name = request.data.get("name", None)

        if ptDocId is None or name is None:
            data = {"code": 201, "data": {}, "msg": "参数错误"}
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)

        try:
            pt = PlanByUserDocument.objects.get(id=ptDocId)
            if name is not None:
                pt.name = name
            pt.save()   
            data = {"code": 200, "data": {}, "msg": "更新成功！", "success": True}
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}

        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    


def deep_copy_model(instance:PlanTemplate, mydate: str, user: User=None):
    # 获取模型类
    new_instance = PlanByUser.objects.create(ctype=instance.ctype, 
                                             yadate=mydate,
                                             user=user,
                                             name=getYuAnName(instance.ctype, mydate),)
    # 保存新实例，生成新的主键
    new_instance.save()
    for node in instance.nodeOutlineList:
        # 处理节点外键字段
        new_instance_node = TemplateNode.objects.create(label=node['label'], description=node['description'], template=node['template'], order=node['order'])
        new_instance.nodes.add(new_instance_node)
    # 保存更新后的新实例
    new_instance.save()
    return new_instance


class YuAnRecomApiPost(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='用户输入预案，动态推荐预览模板',
        operation_description='POST 标准接口',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['text'],
            properties={
                'uid': openapi.Schema(type=openapi.TYPE_INTEGER, description="uid"),
                'text': openapi.Schema(type=openapi.TYPE_STRING, description="text"),
                'date': openapi.Schema(type=openapi.TYPE_STRING, description="date"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        """
            针对用户输入,动态推荐预案摸板
        """
        uid = request.data.get("uid", None)
        text = request.data.get("text", None)
        mydate = request.data.get("date", str(datetime.now().strftime("%Y-%m-%d")))

        # TODO
        pts = PlanTemplate.objects.all()
        all_templates = [
            pt.name for pt in PlanTemplate.objects.all()
        ]
        # template_name = map_input_to_template(text, all_templates)
        # if template_name:
        #     data = model_to_dict(PlanTemplate.objects.get(name=template_name))
        # else:
        #     data = {text}
        try:
            tmpUser = User.objects.get(id=uid)
        except:
            tmpUser = None

        results = recommend_plan(text, plans=pts, user=tmpUser)
        if results:
            target_plan = results[0]
            ptId = target_plan['id']
            ptType = target_plan['ctype']

            try:
                pt = PlanTemplate.objects.get(id=ptId)
            except:
                data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
                bars = YuAnAppResponseSerializer(data=data, many=False)
                bars.is_valid()
                return Response(bars.data, status=status.HTTP_200_OK)
            
            # 深拷贝模型实例
            new_pt = deep_copy_model(pt, mydate, user=tmpUser)
            tmpResult = model_to_dict(new_pt, exclude=["nodes"])
            tmpResult['nodeOutlineList'] = new_pt.nodeOutlineList
            tmpResult['created_at'] = new_pt.created_at.strftime("%Y-%m-%d %H:%M:%S")
            tmpResult['name'] = getYuAnName(ptType, mydate)
            tmpResult['yadate'] = mydate
            data = {"code": 200, "msg": "success", "data": tmpResult, "success": True}
        else:
            data = {"code": 201, "data": {}, "msg": "暂无推荐数据", "success": False}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)

class YuAnRecomPtDetailApiGet(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="获取预案模板列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('id', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        ptId = request.GET.get("id", 1)
        try:
            pt = PlanTemplate.objects.get(id=ptId)
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
            bars = YuAnAppResponseSerializer(data=data, many=False)
            bars.is_valid()
            return Response(bars.data, status=status.HTTP_200_OK)

        tmpResult = model_to_dict(pt, exclude=["nodes", "id"])
        tmpResult['nodeOutlineList'] = pt.nodeOutlineList
        tmpResult['created_at'] = pt.created_at.strftime("%Y-%m-%d %H:%M:%S")
        data = {"code": 200, "msg": "success", "data": tmpResult, "success": True}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)
    


class YuAnUserPtSaveApiPost(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='用户保存预案结构',
        operation_description='POST 标准接口',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['uid'],
            properties={
                'uid': openapi.Schema(type=openapi.TYPE_NUMBER, description="uid"),
                'id': openapi.Schema(type=openapi.TYPE_NUMBER, description="id"),
                'name': openapi.Schema(type=openapi.TYPE_STRING, description="name"),
                'yadate': openapi.Schema(type=openapi.TYPE_STRING, description="yadate"),
                'ctype': openapi.Schema(type=openapi.TYPE_NUMBER, description="ctype"),
                # 'nodeList': openapi.Schema(type=openapi.TYPE_ARRAY,  description="节点列表",  items=openapi.Schema(type=openapi.TYPE_OBJECT, properties={})),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        """
            针对用户输入,动态推荐预案摸板
        """ 
        params = request.data
        print(params)
        uid = params.get("uid", None)
        pid = params.get("id", None)
        name = params.get("name", None)
        yadate = params.get("yadate", str(datetime.now().strftime("%Y-%m-%d")))
        ctype = params.get("ctype", None)
        nodeList = params.get("nodeList", [])
        print("uid:", uid, name, nodeList)
        # TODO
        if pid is None:
            tmpP = PlanByUser.objects.create(name=name, yadate=yadate,ctype=ctype)
        else:
            tmpP = PlanByUser.objects.get(id=pid)

        for nodeDict in nodeList:
            try:
                # 先尝试更新,
                tmpNid = nodeDict.get("id", None)
                tmpN = TemplateNode.objects.get(id=tmpNid)
                for k, v in nodeDict.items():
                    setattr(tmpN, k, v)
                tmpN.save()
            except:
                # 没有则创建
                del nodeDict['id']
                tmpN = TemplateNode.objects.create(**nodeDict)
                tmpP.nodes.add(tmpN)
                
        tmpP.save()
        result = model_to_dict(tmpP, exclude=["nodes"])
        result['nodeOutlineList'] = tmpP.nodeOutlineList
        result['created_at'] = tmpP.created_at.strftime("%Y-%m-%d %H:%M:%S")
        data = {"code": 200, "data": result, "msg": "success"}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)

class YuAnUserPtNodeAddOrUpdateApiPost(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='用户增/修改某个节点',
        operation_description='POST 标准接口',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['id'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_NUMBER, description="text"),
                'nodeId': openapi.Schema(type=openapi.TYPE_NUMBER, description="nodeId"),
                'label': openapi.Schema(type=openapi.TYPE_STRING, description="label"),
                'description': openapi.Schema(type=openapi.TYPE_STRING, description="description"),
                'template': openapi.Schema(type=openapi.TYPE_STRING, description="template"),
                'order': openapi.Schema(type=openapi.TYPE_NUMBER, description="order"),
                'upload_file_flag': openapi.Schema(type=openapi.TYPE_STRING, description="upload_file_flag"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        """
            针对用户输入,动态推荐预案摸板
        """
        pid = request.data.get("id", None)
        nodeId = request.data.get("nodeId", None)
        label = request.data.get("label", None)
        description = request.data.get("description", None)
        template = request.data.get("template", "默认模板")
        order = request.data.get("order", 0)
        upload_file_flag = request.data.get("upload_file_flag", 0)

        try:
            pt = PlanByUser.objects.get(id=pid)
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
            bars = YuAnAppResponseSerializer(data=data, many=False)
            bars.is_valid()
            return Response(bars.data, status=status.HTTP_200_OK)
        
        if nodeId is not None:
            tmpN = TemplateNode.objects.get(id=nodeId)
            if label is not None:
                tmpN.label = label
            if description is not None:
                tmpN.description = description
            if template is not None:
                tmpN.template = template
            if order is not None:
                tmpN.order = order
            if upload_file_flag is not None:
                tmpN.upload_file_flag = upload_file_flag
            tmpN.save()
            data = {"code": 200, "data":{}, "msg": "更新节点内容成功", "success": True}
        else:
            tmpN = TemplateNode.objects.create(label=label, description=description, template=template, order=order, upload_file_flag=upload_file_flag)
            pt.nodes.add(tmpN)
            pt.save()
            data = {"code": 200, "data":{}, "msg": "新建节点成功", "success": True}

        tmpdict = model_to_dict(tmpN, exclude=["result", "wordParagraphs"])
        tmpdict['title'] = tmpN.label
        tmpdict['prompt'] = tmpN.template
        tmpdict['created_at'] = tmpN.created_at.strftime("%Y-%m-%d %H:%M:%S")
        data['data'] = tmpdict
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)


class YuAnUserPtDeleteApiPost(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='用户删除某个节点',
        operation_description='POST 标准接口',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['text'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_NUMBER, description="text"),
                'nodeId': openapi.Schema(type=openapi.TYPE_NUMBER, description="nodeId"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        """
            针对用户输入,动态推荐预案摸板
        """
        pid = request.data.get("id", None)
        nodeId = request.data.get("nodeId", None)

        try:
            pt = PlanByUser.objects.get(id=pid)
            tmpN = TemplateNode.objects.get(id=nodeId)
            pt.nodes.remove(tmpN)
            tmpN.delete()
            data = {"code": 200, "data": {}, "msg": "success", "success": True}
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    

class YuAnUserPtDetailApiGet(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="获取用户预案列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('id', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        
        ptId = request.GET.get("id", 1)
        try:
            pt = PlanByUser.objects.get(id=ptId)
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
            bars = YuAnAppResponseSerializer(data=data, many=False)
            bars.is_valid()
            return Response(bars.data, status=status.HTTP_200_OK)

        tmpResult = model_to_dict(pt, exclude=["nodes"])
        tmpResult['nodeOutlineList'] = pt.nodeOutlineList
        tmpResult['created_at'] = pt.created_at.strftime("%Y-%m-%d %H:%M:%S")
        data = {"code": 200, "msg": "success", "data": tmpResult, "success": True}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)


class YuAnUserPtListApiGet(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="[生成预案列表]  获取用户预案列表",
        # 接口参数 GET请求参数
        manual_parameters=[
            openapi.Parameter('uid', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
            openapi.Parameter('keyword', openapi.IN_QUERY, type=openapi.TYPE_STRING),
            openapi.Parameter('page', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
            openapi.Parameter('pageSize', openapi.IN_QUERY, type=openapi.TYPE_NUMBER),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        uid = request.GET.get("uid", None)
        keyword = request.GET.get("keyword", None)
        page = request.GET.get("page", 1)
        pageSize = request.GET.get("pageSize", 5)
        
        tmpUser = User.objects.get(id=uid)
        if tmpUser is None:
            YuAnUserPtListApiGet.queryset = PlanByUser.objects
        else:
            YuAnUserPtListApiGet.queryset = PlanByUser.objects.filter(user=tmpUser)

        if keyword:
            YuAnUserPtListApiGet.queryset = YuAnUserPtListApiGet.queryset.filter(name__icontains=keyword).order_by('-created_at').all()
        else:
            YuAnUserPtListApiGet.queryset = YuAnUserPtListApiGet.queryset.order_by('-created_at').all()

        paginator = Paginator(YuAnUserPtListApiGet.queryset, pageSize)
        try:
            ptList = paginator.page(page)
        except PageNotAnInteger:
            ptList = paginator.page(1)
        except EmptyPage:
            ptList = paginator.page(paginator.num_pages)

        results = []
        for tmp in ptList:
            tmpdict = model_to_dict(tmp, exclude=["plan", "nodes"])
            tmpdict['created_at'] = tmp.created_at.strftime("%Y-%m-%d %H:%M:%S")
            results.append(tmpdict)
        print(results)
        data = {"code": 200, "msg": "success", "data": results, "success": True, "total": paginator.count, "page": paginator.num_pages, "pageSize": paginator.per_page}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)
    

class RecentlyYuAnUserPtApiGet(mixins.ListModelMixin,
                           mixins.CreateModelMixin,
                           generics.GenericAPIView):
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_description='GET ///',
        operation_summary="[获取最新预案]  获取用户最近编辑的",
        # 接口参数 GET请求参数
        manual_parameters=[
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api'])
    def get(self, request, *args, **kwargs):
        uid = request.GET.get("uid", None)
        try:
            tmpUser = User.objects.get(id=uid)
        except:
            tmpUser = None
            
        if tmpUser is None:
            YuAnUserPtListApiGet.queryset = PlanByUser.objects
        else:
            YuAnUserPtListApiGet.queryset = PlanByUser.objects.filter(user=tmpUser)
        recentUserPt = YuAnUserPtListApiGet.queryset.order_by('-created_at').first()
        if recentUserPt is None:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
            bars = YuAnAppResponseSerializer(data=data, many=False)
            bars.is_valid()
            return Response(bars.data, status=status.HTTP_200_OK)
        result = model_to_dict(recentUserPt, exclude=["plan", "nodes"])
        result['nodeOutlineList'] = recentUserPt.nodeOutlineList
        result['created_at'] = recentUserPt.created_at.strftime("%Y-%m-%d %H:%M:%S")
        data = {"code": 200, "msg": "success", "data": result, "success": True}
        bars = YuAnAppResponseSerializer(data=data, many=False)
        bars.is_valid()
        return Response(bars.data, status=status.HTTP_200_OK)
    
    
class YuAnUserPlanDeleteApiPost(generics.GenericAPIView):
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='用户删除预案',
        operation_description='POST 标准接口',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['text'],
            properties={
                'id': openapi.Schema(type=openapi.TYPE_NUMBER, description="text"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        """
            针对用户输入,动态推荐预案摸板
        """
        pid = request.data.get("id", None)
        try:
            pt = PlanByUser.objects.get(id=pid)
            pt.nodes.clear()
            pt.delete()
            data = {"code": 200, "data": {}, "msg": "success", "success": True}
        except:
            data = {"code": 202, "data": {}, "msg": "系统不存在该数据", "success": False}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)
    
class DDFAUploadApiPost(generics.GenericAPIView):
    parser_classes = (FormParser, MultiPartParser)
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='POST 调度方案单上传',
        operation_description='POST /ddfadUpload',
        manual_parameters=[
            openapi.Parameter(
                name='myfile',
                in_=openapi.IN_FORM,
                description='调度方案单',
                type=openapi.TYPE_FILE
            ),
            openapi.Parameter('mydate', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='方案日期',),
            openapi.Parameter('mytype', openapi.IN_QUERY, type=openapi.TYPE_STRING, description='方案类型',),
        ],
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags=['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        print(request.FILES, flush=True)
        myFile = request.FILES.get("myfile", None)
        myType = request.GET.get("mytype", 0)
        myDate = request.GET.get("mydate", str(datetime.now().strftime("%Y-%m-%d")))
        if not myFile:
            krrs = YuAnAppResponseSerializer(data={"code": 200, "msg": "No document provided.!"}, many=False)
            krrs.is_valid()
            return Response(krrs.data, status=status.HTTP_400_BAD_REQUEST)
        
        tmproot = os.path.join("data", "yuan_data", str(myType), "ddfad")
        if not os.path.exists(tmproot):
            os.makedirs(tmproot)
        df_path = os.path.join(tmproot, f"{myDate}.xlsx")
        f = open(df_path, "wb+")
        # 分块写入文件
        for chunk in myFile.chunks():
            f.write(chunk)
        f.close()
        print("调度方案单写入:", df_path, flush=True)
        yautils.plot_save_html(df_path, business_type=myType, myDate=myDate)
        print("绘图完成")
        krrs = YuAnAppResponseSerializer(data={"code": 200, "msg": "上传调度方案单成功", "success": True}, many=False)
        krrs.is_valid()
        return Response(krrs.data, status=status.HTTP_200_OK)


def downloadPlan(request):
    docId = request.GET.get('docId')
    tmpDoc = PlanByUserDocument.objects.get(id=docId)
    print("downloadPlan-->", tmpDoc)
    file_path = os.path.join("media", "plans", f"{tmpDoc.name}.docx")
    file_handle = open(file_path, 'rb')
    # 创建 FileResponse 实例
    response = FileResponse(file_handle)
    response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # 设置文件名，这将影响下载对话框中显示的文件名
    response['Content-Disposition'] = f'attachment; filename="{tmpDoc.name}.docx"'

    return response



class LLMYuAnTaskApiView(generics.GenericAPIView):

    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer

    @swagger_auto_schema(
        operation_summary='[可用] 预案生成任务',
        operation_description='POST /llmYuAnTask',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=["user_id", "ptid"],
            properties={
                'user_id': openapi.Schema(type=openapi.TYPE_INTEGER, description="创建作者ID"),
                'name': openapi.Schema(type=openapi.TYPE_STRING, description="任务名称"),
                'desc': openapi.Schema(type=openapi.TYPE_STRING, description="任务描述"),
                'ptid': openapi.Schema(type=openapi.TYPE_INTEGER, description="预案ID"),
                'taskType': openapi.Schema(type=openapi.TYPE_INTEGER,
                                           description="生产类型 0(知识导入)|1(自动生成)|2(全文生产)|3(图谱导入)|4(图谱自动生产)|5(预案生成)"),
            },
        ),
        responses={
            200: YuAnAppResponseSerializer(many=False),
            400: "请求失败",
        },
        tags = ['ya_api']
    )
    @csrf_exempt
    def post(self, request, *args, **krgs):
        data = {"code": 200 }
        name = request.data.get("name", None)
        taskType = int(request.data.get("taskType", 5))
        desc = request.data.get("desc", "预案生成任务")
        user_id = request.data.get("user_id", 1)
        planId = request.data.get("ptid", None)

        if name is None or user_id is None or not planId:
            data['code'] = 201
            data['msg'] = '请求参数错误, 缺少参数！！！'
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data,  status=status.HTTP_200_OK)
        else:
            
            try:
                tmpuser = User.objects.get(id=user_id)
            except:
                data = {"code": 202, "msg": "用户ID不存在！！！" }
                serializers = YuAnAppResponseSerializer(data=data, many=False)
                serializers.is_valid()
                return Response(serializers.data, status=status.HTTP_200_OK)
            

            tmptask = KgProductTask()
            tmptask.created_at = datetime.now()
            tmptask.updated_at = datetime.now()
            tmptask.name = name
            tmptask.desc = desc
            tmptask.task_type = taskType
            
            tmptask.kg_user_id = tmpuser
            tmptask.task_status = 0  #### 0 未执行,  1 任务开启, 执行数据装载, 2 数据装载完成, 3 比对任务开启, 4 比对任务完成, 5 最终任务完成, -1 任务失败
            tmptask.save()

            from yunheKGPro.celery import LLMYuAnProdTask
            param_dict = model_to_dict(tmptask, exclude=['kg_doc_ids'])
            param_dict['planId'] = planId
            param_dict['task_id'] = tmptask.id
            result = LLMYuAnProdTask.delay(param_dict)
            # 记录 celery id 入库 
            subtask, tmpbool = KgTask.objects.get_or_create(kg_prod_task_id=tmptask, task_step=1, celery_id=result.id)
            if tmpbool:
                print("ProductTask 任务创建成功...", subtask)

            data = {"code": 200, "msg": "预案生产任务新建成功", "success": True, "data": model_to_dict(subtask)}
            data['data'] = model_to_dict(subtask)
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)



class LLMYuAnTaskStatusApiView(mixins.ListModelMixin,
                  mixins.CreateModelMixin,
                  generics.GenericAPIView):
    
    serializer_class = YuAnAppResponseSerializer
    @swagger_auto_schema(
            operation_description='GET /llmYuAnTaskStatus',
            operation_summary="异步任务状态",
            # 接口参数 GET请求参数
            manual_parameters=[
                # 声明参数
                openapi.Parameter(
                    "taskid", 
                    openapi.IN_QUERY, 
                    description="任务ID", 
                    type=openapi.TYPE_INTEGER
                ),
            ],
            responses={
                200: YuAnAppResponseSerializer(many=False),
                400: "请求失败",
            },
            tags = ['ya_api'])
    def get(self, request, *args, **kwargs):
        taskid = request.GET.get("taskid", None)
        if taskid is None:
            data = {"code": 201, "msg": "参数错误！！！" }
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        
        tmptask = KgTask.objects.filter(kg_prod_task_id=taskid).first()
        if tmptask:
            from celery.result import AsyncResult
            celery_task = AsyncResult(tmptask.celery_id)
            celery_status = celery_task.status
        else:
            celery_status = 'TASK_NOT_FOUND'
        data = {"code": 200, "msg": "", "status":  celery_status}
        serializers = YuAnAppResponseSerializer(data=data, many=False)
        serializers.is_valid()
        return Response(serializers.data, status=status.HTTP_200_OK)




class ResultFromTaskApiView(mixins.ListModelMixin,
                  mixins.CreateModelMixin,
                  generics.GenericAPIView):
    
    authentication_classes = (CsrfExemptSessionAuthentication, BasicAuthentication)
    serializer_class = YuAnAppResponseSerializer
    
    
    @swagger_auto_schema(
            operation_description='GET /resultFromTask',
            operation_summary="获取生成的结果",
            # 接口参数 GET请求参数
            request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['task_id'],
            properties={
                'task_id': openapi.Schema(type=openapi.TYPE_INTEGER, description="任务ID"),
            },
        ),
            responses={
                200: YuAnAppResponseSerializer(many=False),
                400: "请求失败",
            },
            tags = ['ya_api'])
    def post(self, request, *args, **kwargs):

        # 根据不同的任务，启动不同的装载数据方式
        task_id = request.data.get("task_id", None)
        tmptask = KgTask.objects.get(id=task_id)
        if tmptask is None:
            data = {"code": 202, "msg": "暂无生产子任务失败" }
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        
        from celery.result import AsyncResult
        import pprint
        res = AsyncResult(tmptask.celery_id) # 参数为task id
        if not res:
            data = {"code": 200, "msg": "任务ID结果不存在" }
            serializers = YuAnAppResponseSerializer(data=data, many=False)
            serializers.is_valid()
            return Response(serializers.data, status=status.HTTP_200_OK)
        print("res.result:", res.result)
        resultSet = res.result['data']
        serializers = YuAnAppResponseSerializer(data=resultSet, many=False)
        serializers.is_valid()
        return Response(serializers.data,  status=status.HTTP_200_OK)