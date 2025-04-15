from io import BytesIO
from docx.shared import Pt  
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import re
import json
import pandas as pd
from openpyxl import load_workbook
from openpyxl.cell import MergedCell
from docx.shared import Cm, Mm, Inches

# 定义一个函数来设置单元格边框

from docx import Document
from docx.shared import Pt, Inches

def set_landscape(doc):
    sections = doc.sections
    if not sections:
        section = doc.add_section()
    else:
        section = sections[0]
    
    # 获取当前的页面设置
    page_height = section.page_height
    page_width = section.page_width
    
    # 交换宽度和高度以设置为横向
    section.page_height = page_width
    section.page_width = page_height
    
    # 设置页边距
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    # 如果有奇偶页边距，也需要设置
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "color": "#FF0000", "val": "single"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "color": "#0000FF", "val": "double"},
        end={"sz": 12, "color": "#FF00FF", "val": "single"}
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def writeTitle2Word(title, doc):
    # 添加一个段落
    titlePara = doc.add_paragraph(title)
    # 设置字体名称和大小
    run = titlePara.runs[0]
    # run.font.name = 'Arial'
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(24)
    # 设置字体加粗
    run.font.bold = True
    # 设置字体斜体
    # run.font.italic = True
    titlePara.alignment = WD_ALIGN_PARAGRAPH.CENTER


import base64
from docx import Document
from docx.shared import Inches

#bianHaos = ['一','二','三','四','五','六','七','八','九','十', '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']
bianHaos = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13', '14', '15', '16', '17', '18', '19', '20']


def writeParagraphs2Word(i, node: dict, doc):
    # 添加标题（黑体，数字强制 Times New Roman）
    subTitle = node['label']
    heading = doc.add_heading(f"{bianHaos[i]}、 {subTitle}", level=1)
    for run in heading.runs:
        run.font.name = 'Times New Roman'  # 数字和英文用 Times New Roman
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 中文用黑体
        run.bold = True
        run.font.size = Pt(16)

    # 处理正文段落
    for para in node['paraglist']:
        if para['ctype'] == 1:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(para['content'])
            run.font.name = 'Times New Roman'  # 数字和英文用 Times New Roman
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 中文用仿宋
            run.font.size = Pt(16)

            if len(para['content'].strip()) < 10:
                run.font.name = 'Times New Roman'  # 数字和英文用 Times New Roman
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')  # 中文用楷体
                run.bold = True
                run.font.size = Pt(16) # 三号字体（Word 三号 = 16磅）
        if para['ctype'] == 2:
                imgdata = base64.b64decode(para['content'])
                img_stream = BytesIO(imgdata)
                # 添加段落用于存放图片
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中
                run = paragraph.add_run()
                run.add_picture(img_stream, width=Inches(5.0))
                # 将图片添加到Word文档
                # 注意：add_picture 方法通常期望一个文件路径，但在这里我们使用BytesIO对象作为文件流
                #doc.add_picture(img_stream, width=Inches(5.0))
        if para['ctype'] == 3:
            # 表格
            json_str = json.loads(para['content'])
            df = pd.read_json(json_str)  
            # 将DataFrame添加到Word文档中
            table = doc.add_table(rows=1, cols=len(df.columns))

            # 设置表格的宽度
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl_width = OxmlElement('w:tblW')
            tbl_width.set(qn('w:type'), 'auto')
            tbl_width.set(qn('w:w'), str(5000))
            table._element.tblPr.append(tbl_width)

            # 添加表头
            for j, col in enumerate(df.columns):
                cell = table.cell(0, j)
                cell.text = str(col)
                # 设置表头单元格边框
                cellparagraph = cell.paragraphs[0]
                cellparagraph_format = cellparagraph.paragraph_format
                cellparagraph_format.space_after = Pt(0)
                set_cell_border(cell, top={"sz": 12, "color": "#000000", "val": "single"},
                                bottom={"sz": 12, "color": "#000000", "val": "single"},
                                start={"sz": 12, "color": "#000000", "val": "single"},
                                end={"sz": 12, "color": "#000000", "val": "single"})
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                        run.font.size = Pt(10.5)
            # 添加DataFrame的内容
            for i, row in df.iterrows():
                row_cells = table.add_row().cells
                for j, val in enumerate(row):
                    cell = row_cells[j]
                    cell.text = str(val)
                    # 设置数据单元格边框
                    cellparagraph = cell.paragraphs[0]
                    cellparagraph_format = cellparagraph.paragraph_format
                    cellparagraph_format.space_after = Pt(0)
                    set_cell_border(cell, top={"sz": 12, "color": "#000000", "val": "single"},
                                    bottom={"sz": 12, "color": "#000000", "val": "single"},
                                    start={"sz": 12, "color": "#000000", "val": "single"},
                                    end={"sz": 12, "color": "#000000", "val": "single"})
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                            run.font.size = Pt(10.5)
        if para['ctype'] == 4:
            pass




def ddfa_html_table(ddfa_excel='data/ddfa/tmp.xlsx'):
    wb = load_workbook(filename=ddfa_excel)
    sheet = wb['Sheet1']
    cell_dic = {}  # 用于储存 所有合并单元格的左上单元格对象
    col_width = {}  # 用于储存 所有列的列宽,px
    row_height = {}  # 用于储存 所有列的行高,px
    # 查询列宽
    for col in sheet.columns:
        pat = r"[A-Z]+"
        pat = re.compile(pat)
        colname = pat.findall(col[0].coordinate)[0]  # 分离字母 和数字，取出列字母名称
        px = round(sheet.column_dimensions[colname].width * 5)  # 读出列宽换算为像素
        col_width[colname] = px
    # 查询行高
    for row in sheet.rows:
        pat = r"[A-Z]+(\d+)"
        pat = re.compile(pat)
        rowid = int(pat.findall(row[0].coordinate)[0])  # 分离字母 和数字，取出行数字序号
        px = sheet.row_dimensions[rowid].height  # 读出行高换算为像素
        if px == None:
            px = 13.5
        row_height[str(rowid)] = px
    # 找出所有合并区域的行高，列宽，向右合并距离，向下合并距离
    for merged_range in sheet.merged_cells.ranges:
        now_width = 0  # 定义列宽
        now_height = 0  # 定义行高
        for i in range(merged_range.min_col, merged_range.max_col + 1):
            coord = sheet.cell(row=1, column=i).coordinate  # 位置标识，例如：A1
            pat = r"[A-Z]+"
            pat = re.compile(pat)
            colname = pat.findall(coord)[0]  # 分离字母 和数字，取出列字母名称
            now_width = now_width + col_width[colname]
        for i in range(merged_range.min_row, merged_range.max_row + 1):
            coord = sheet.cell(row=i, column=1).coordinate  # 位置标识，例如：A1
            pat = r"[A-Z]+(\d+)"
            pat = re.compile(pat)
            colindex = pat.findall(coord)[0]  # 分离字母 和数字，取出列数字名称
            now_height = now_height + row_height[colindex]

        now_width = int(now_width)  # 合并单元格列宽（所有子单元格相加）
        now_height = int(now_height)  # 合并单元格行高（所有子单元格相加）

        cell = sheet.cell(row=merged_range.min_row, column=merged_range.min_col)  # 选择合并区域左上单元格
        colspan = merged_range.max_col - merged_range.min_col + 1  # 向右合并长度
        rowspan = merged_range.max_row - merged_range.min_row + 1  # 向下合并长度
        cell_dic[cell] = (now_height, now_width, colspan, rowspan)

    ddfa_html = '''<table cellspacing="0" class="Table" style="background:white; border-collapse:collapse; margin-left:64px;">'''
    # 开始写入数据到table标签
    for row in sheet.rows:
        tr = '''<tr>'''
        for cell in row:
            td = ""
            if cell in cell_dic:  # 判断是否为合并单元格左上单元格
                if cell.value == None:
                    text = ''
                else:
                    text = cell.value

                if cell.alignment.vertical != None:
                    vertical = f'''vertical-align: {cell.alignment.vertical};'''  # 水平位置
                else:
                    vertical = ''
                if cell.alignment.horizontal != None:
                    horizontal = f'''text-align: {cell.alignment.horizontal};'''  # 垂直位置
                else:
                    horizontal = ''

                font_size = str(int(cell.font.size) + 3)  # 字体大小
                font_weight = '700' if cell.font.b else '400'  # 字体是否加粗
                style = f'''"color: rgb(0, 0, 0); font-size: {font_size}px; font-weight: {font_weight}; font-style: normal;{vertical}{horizontal}"'''
                td = f'''<td height="{cell_dic[cell][0]}" width="{cell_dic[cell][1]}" colspan="{cell_dic[cell][2]}" rowspan="{cell_dic[cell][3]}" style={style}>{text}</td>'''
            else:
                if not isinstance(cell, MergedCell):  # 判断该单元格是否为合并单元格
                    if cell.alignment.vertical != None:
                        vertical = f'''vertical-align: {cell.alignment.vertical};'''  # 水平位置
                    else:
                        vertical = ''
                    if cell.alignment.horizontal != None:
                        horizontal = f'''text-align: {cell.alignment.horizontal};'''  # 垂直位置
                    else:
                        horizontal = ''
                    pat = r"([A-Z]+)(\d+)"
                    pat = re.compile(pat)
                    cell_name = pat.findall(cell.coordinate)[0][0]
                    cell_index = pat.findall(cell.coordinate)[0][1]
                    font_size = str(int(cell.font.size) + 3)  # 字体大小
                    font_weight = '700' if cell.font.b else '400'  # 字体是否加粗
                    style = f'''"background-color:white; border-bottom:1px solid black; border-left:1px solid black; border-right:1px solid black; border-top:1px solid black;"'''
                    if cell.value != None:
                        if isinstance(cell.value, float):
                            td = f'''<td height="{row_height[cell_index]}" width="{col_width[cell_name]}" style={style} >{round(cell.value, 2)}</td>'''
                        elif isinstance(cell.value, datetime.datetime):
                            td = f'''<td height="{row_height[cell_index]}" width="{col_width[cell_name]}" style={style} >{cell.value.strftime('%Y-%m-%d %H:%M:%S')}</td>'''
                        else:
                            td = f'''<td height="{row_height[cell_index]}" width="{col_width[cell_name]}" style={style} >{cell.value}</td>'''
                    else:
                        td = f'''<td height="{row_height[cell_index]}" width="{col_width[cell_name]}"></td>'''
            tr = tr + td
        tr = tr + '''</tr>'''
        ddfa_html = ddfa_html + tr
    html = ddfa_html + '''</table>'''
    return html
